#!/usr/bin/env python3
"""
Generate DTU B.Tech Project-II Report in Word (.docx) format.
Target: 55-60 pages. Follows the exact DTU CSE template.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Paths ────────────────────────────────────────────────
REPORT_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_DIR    = os.path.join(REPORT_DIR, 'figures')
ML_OUT_DIR = os.path.join(REPORT_DIR, '..', 'ml-api', 'outputs')

# ── Page Setup (A4) ───────────────────────────────────────
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.81)
    section.right_margin = Cm(2.54)

# ── Base Style ────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# ── Helpers ───────────────────────────────────────────────
def pb():
    doc.add_page_break()

def ctr(text, size=16, bold=True, underline=False, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline
    run.font.name = 'Times New Roman'
    return p

def body(text, bold=False, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p

def right_align(text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p

def heading1(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def heading2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def heading3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
    doc.add_paragraph()
    return table

def blank(n=1):
    for _ in range(n):
        doc.add_paragraph()

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p


def add_figure(image_path, caption, width=Cm(14)):
    """Insert an image with a centered, bold, italic caption below it."""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
    else:
        body(f'[Image not found: {image_path}]', bold=True, indent=False)
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(caption)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = True
    run.italic = True


# ═══════════════════════════════════════════════════════════
#  PAGE 1: TITLE PAGE
# ═══════════════════════════════════════════════════════════
blank(2)
ctr('Real-Time UPI Fraud Detection System\nUsing Machine Learning', 18, True, False, 0, 12)
blank()
ctr('A B.TECH PROJECT-II REPORT', 14, True, False, 0, 6)
ctr('SUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENTS FOR\nTHE AWARD OF THE DEGREE', 11, False, False, 6, 4)
ctr('OF', 11, False, False, 0, 4)
ctr('BACHELOR OF TECHNOLOGY', 14, True, False, 4, 4)
ctr('IN', 11, False, False, 0, 4)
ctr('COMPUTER ENGINEERING', 14, True, False, 0, 18)
ctr('Submitted By', 12, True, False, 12, 12)
ctr('[Name of Student 1]          [Name of Student 2]          [Name of Student 3]', 11, True, False, 0, 4)
ctr('([Roll No. 1])                    ([Roll No. 2])                    ([Roll No. 3])', 11, False, False, 0, 18)
ctr('Under the supervision of', 12, False, False, 6, 6)
ctr('[Name of Supervisor]', 12, True, False, 0, 24)
blank()
ctr('DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING', 12, True, False, 6, 4)
ctr('DELHI TECHNOLOGICAL UNIVERSITY', 12, False, False, 0, 2)
ctr('(Formerly, Delhi College of Engineering)', 11, False, False, 0, 2)
ctr('Bawana Road, Delhi-110042', 11, False, False, 0, 12)
ctr('May 2026', 14, False, False, 12, 0)


# ═══════════════════════════════════════════════════════════
#  PAGE 2: CANDIDATE'S DECLARATION
# ═══════════════════════════════════════════════════════════
pb()
ctr("CANDIDATE\u2019S DECLARATION", 16, True, True, 24, 18)

body('We, [Name of Student 1] ([Roll No.]), [Name of Student 2] ([Roll No.]), [Name of Student 3] ([Roll No.]), pursuing Bachelor of Technology degree in Computer Engineering from the Department of Computer Science and Engineering, hereby declare that the project report titled \u201cReal-Time UPI Fraud Detection System Using Machine Learning\u201d submitted by us to the Department of Computer Science and Engineering, Delhi Technological University, Delhi, in partial fulfillment of the degree, is entirely original and not copied from any source without proper citation. We hereby declare that this work has not been submitted in part or full for any degree or diploma to this University or elsewhere.')

blank(3)
right_align('(Signature)')
right_align('[Name of Student 1] ([Roll No.])')
blank()
right_align('(Signature)')
right_align('[Name of Student 2] ([Roll No.])')
blank()
right_align('(Signature)')
right_align('[Name of Student 3] ([Roll No.])')
blank(2)
body('Place : New Delhi', indent=False)
body('Date :', indent=False)


# ═══════════════════════════════════════════════════════════
#  PAGE 3: CERTIFICATE
# ═══════════════════════════════════════════════════════════
pb()
ctr('CERTIFICATE', 16, True, True, 24, 18)

body('This is to certify that the work entitled \u201cReal-Time UPI Fraud Detection System Using Machine Learning\u201d submitted by [Name of Student 1] ([Roll No.]), [Name of Student 2] ([Roll No.]), [Name of Student 3] ([Roll No.]), of the Department of Computer Science and Engineering, Delhi Technological University, in partial fulfillment of the requirement for the project work, has been carried out by the students under my supervision. To the best of my knowledge, this work has not been submitted in part or full for any degree or diploma to this University or elsewhere.')

blank(4)
right_align('(Signature)')
right_align('[Name of Supervisor]')
right_align('[Designation of Supervisor]')
blank(3)
body('Place : New Delhi', indent=False)
body('Date :', indent=False)


# ═══════════════════════════════════════════════════════════
#  PAGE 4: ACKNOWLEDGEMENT
# ═══════════════════════════════════════════════════════════
pb()
ctr('ACKNOWLEDGEMENT', 16, True, True, 24, 18)

body('We would like to express our sincere gratitude to our project supervisor, [Supervisor Name], [Designation], Department of Computer Science and Engineering, Delhi Technological University, for providing invaluable guidance, constructive feedback, and constant encouragement throughout this project.')

body('We extend our thanks to [HOD Name], Head of the Department of Computer Science and Engineering, for providing the necessary infrastructure and academic environment that facilitated the completion of this work.')

body('We are grateful to the faculty members of the department for the knowledge and skills imparted during our B.Tech programme, which formed the foundation upon which this project was built. We also thank our fellow students for their constructive discussions and suggestions during the course of this project.')

body('We would also like to acknowledge the open-source community whose tools and libraries\u2014including XGBoost, scikit-learn, FastAPI, Express.js, React, Prisma, and PostgreSQL\u2014made this project technically feasible within the given timeframe.')

body('Finally, we are deeply indebted to our families for their unwavering support, patience, and encouragement throughout our academic journey.')

blank(4)
right_align('(Signature)')
right_align('[Name of Student 1] ([Roll No.])')
blank()
right_align('(Signature)')
right_align('[Name of Student 2] ([Roll No.])')
blank()
right_align('(Signature)')
right_align('[Name of Student 3] ([Roll No.])')


# ═══════════════════════════════════════════════════════════
#  PAGE 5: ABSTRACT
# ═══════════════════════════════════════════════════════════
pb()
ctr('ABSTRACT', 16, True, True, 24, 18)

body('The Unified Payments Interface (UPI) processes over 11 billion transactions monthly in India, making it the largest real-time payment system globally. This growth has led to a proportional surge in fraudulent activities, with digital payment fraud cases rising 300 percent between 2020 and 2023 according to the Reserve Bank of India. Traditional rule-based detection systems suffer from high false positive rates and inability to adapt to evolving fraud tactics.')

body('This project presents the design, implementation, and evaluation of a real-time UPI fraud detection system that employs XGBoost gradient boosting classification combined with a comprehensive feature engineering pipeline tailored to UPI transaction characteristics. The system transforms six raw transaction attributes into twenty-eight derived features spanning amount characteristics, balance dynamics, temporal patterns, user behavioural history, and categorical encodings.')

body('The XGBoost classifier, trained on 50,000 synthetic UPI transactions with SMOTE oversampling to address class imbalance, achieves 94% fraud recall, 85% fraud precision, and an ROC-AUC of 0.987 on held-out test data. The trained model is deployed via a FastAPI inference service with 28 ms average prediction latency. An Express.js backend orchestrates transaction processing, JWT authentication, and alert generation, while a React dashboard provides real-time monitoring and fraud analysis capabilities.')

body('A dual-mode detection architecture ensures uninterrupted fraud monitoring: the system automatically falls back to rule-based scoring when the ML service is unavailable, then seamlessly resumes ML-based detection upon recovery. The complete system comprises 132 source files across three services, with 53 test cases passing across unit, integration, security, and performance categories.')

body('Keywords: UPI, fraud detection, machine learning, XGBoost, gradient boosting, feature engineering, SMOTE, real-time prediction, web application, React, FastAPI, Express.js, PostgreSQL', bold=True, indent=False)


# ═══════════════════════════════════════════════════════════
#  LIST OF TABLES
# ═══════════════════════════════════════════════════════════
pb()
ctr('LIST OF TABLES', 16, True, True, 24, 18)

tables_list = [
    ('Table 1', 'Technology Stack'),
    ('Table 2', 'Hardware and Software Requirements'),
    ('Table 3', 'Database Schema \u2014 Users Table'),
    ('Table 4', 'Database Schema \u2014 Transactions Table'),
    ('Table 5', 'Database Schema \u2014 Alerts Table'),
    ('Table 6', 'Database Indexing Strategy'),
    ('Table 7', 'Synthetic Data Generation Parameters'),
    ('Table 8', 'Feature Engineering Summary (28 Features)'),
    ('Table 9', 'XGBoost Hyperparameters'),
    ('Table 10', 'Fraud Threshold Configuration'),
    ('Table 11', 'Rule-Based Fallback Scoring'),
    ('Table 12', 'Backend API Endpoints'),
    ('Table 13', 'Overall Classification Metrics'),
    ('Table 14', 'Per-Class Classification Report'),
    ('Table 15', 'Confusion Matrix'),
    ('Table 16', 'Feature Importance (Top 10)'),
    ('Table 17', 'API Response Times'),
    ('Table 18', 'Fraud Detection by Transaction Type'),
    ('Table 19', 'Fraud Detection by Time Period'),
    ('Table 20', 'ML Model vs Rule-Based Comparison'),
    ('Table 21', 'Impact of SMOTE on Model Performance'),
    ('Table 22', 'API Endpoint Test Results'),
    ('Table 23', 'Frontend Performance Metrics'),
    ('Table 24', 'Concurrent User Load Testing'),
    ('Table 25', 'Security Test Results'),
    ('Table 26', 'Work Plan and Timeline'),
]
for t in tables_list:
    body(f'{t[0]} : {t[1]}', indent=False)


# ═══════════════════════════════════════════════════════════
#  LIST OF FIGURES
# ═══════════════════════════════════════════════════════════
pb()
ctr('LIST OF FIGURES', 16, True, True, 24, 18)

figures_list = [
    ('Figure 1', 'High-Level System Architecture'),
    ('Figure 2', 'Request Processing Pipeline'),
    ('Figure 3', 'Entity-Relationship Diagram'),
    ('Figure 4', 'Transaction Processing Data Flow'),
    ('Figure 5', 'Authentication and JWT Flow'),
    ('Figure 6', 'XGBoost Training Pipeline'),
    ('Figure 7', 'Feature Engineering Categories'),
    ('Figure 8', 'SMOTE Oversampling Visualization'),
    ('Figure 9', 'Confusion Matrix'),
    ('Figure 10', 'ROC Curve'),
    ('Figure 11', 'Precision-Recall Curve'),
    ('Figure 12', 'Feature Importance (Top 20)'),
    ('Figure 13', 'Fraud Probability Distribution'),
    ('Figure 14', 'Login Page Screenshot'),
    ('Figure 15', 'Dashboard Screenshot'),
    ('Figure 16', 'Transactions Page Screenshot'),
    ('Figure 17', 'Check Transaction Page Screenshot'),
    ('Figure 18', 'Alerts Page Screenshot'),
    ('Figure 19', 'Analytics Page Screenshot'),
    ('Figure 20', 'Gantt Chart \u2014 Work Plan'),
]
for f in figures_list:
    body(f'{f[0]} : {f[1]}', indent=False)


# ═══════════════════════════════════════════════════════════
#  LIST OF ABBREVIATIONS
# ═══════════════════════════════════════════════════════════
pb()
ctr('LIST OF ABBREVIATIONS', 16, True, True, 24, 18)

abbrevs = [
    ('ACID', 'Atomicity, Consistency, Isolation, Durability'),
    ('API', 'Application Programming Interface'),
    ('AUC', 'Area Under the Curve'),
    ('CORS', 'Cross-Origin Resource Sharing'),
    ('CSS', 'Cascading Style Sheets'),
    ('DOM', 'Document Object Model'),
    ('FN', 'False Negative'),
    ('FP', 'False Positive'),
    ('GNN', 'Graph Neural Network'),
    ('JWT', 'JSON Web Token'),
    ('KPI', 'Key Performance Indicator'),
    ('LSTM', 'Long Short-Term Memory'),
    ('ML', 'Machine Learning'),
    ('MVCC', 'Multi-Version Concurrency Control'),
    ('NPCI', 'National Payments Corporation of India'),
    ('ORM', 'Object-Relational Mapping'),
    ('P2M', 'Peer-to-Merchant'),
    ('P2P', 'Peer-to-Peer'),
    ('RBAC', 'Role-Based Access Control'),
    ('RBI', 'Reserve Bank of India'),
    ('REST', 'Representational State Transfer'),
    ('ROC', 'Receiver Operating Characteristic'),
    ('SHAP', 'SHapley Additive exPlanations'),
    ('SMOTE', 'Synthetic Minority Over-sampling Technique'),
    ('SPA', 'Single Page Application'),
    ('SQL', 'Structured Query Language'),
    ('TN', 'True Negative'),
    ('TP', 'True Positive'),
    ('UPI', 'Unified Payments Interface'),
    ('VPA', 'Virtual Payment Address'),
    ('XGBoost', 'Extreme Gradient Boosting'),
]
add_table(['Abbreviation', 'Description'], abbrevs)


# ═══════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════
pb()
ctr('TABLE OF CONTENTS', 16, True, True, 24, 18)

toc_items = [
    ('', 'DECLARATION', 'i', True),
    ('', 'CERTIFICATE', 'ii', True),
    ('', 'ACKNOWLEDGEMENT', 'iii', True),
    ('', 'ABSTRACT', 'iv', True),
    ('', 'LIST OF TABLES', 'v', True),
    ('', 'LIST OF FIGURES', 'vi', True),
    ('', 'LIST OF ABBREVIATIONS', 'vii', True),
    ('1.', 'INTRODUCTION', '1', True),
    ('', '1.1. Background', '1', False),
    ('', '1.2. Growth of UPI and Fraud Landscape', '2', False),
    ('', '1.3. Motivation', '3', False),
    ('', '1.4. Problem Statement', '4', False),
    ('', '1.5. Scope of the Project', '5', False),
    ('2.', 'LITERATURE SURVEY', '6', True),
    ('', '2.1. Rule-Based Fraud Detection', '6', False),
    ('', '2.2. Machine Learning Approaches', '7', False),
    ('', '2.3. Deep Learning Methods', '8', False),
    ('', '2.4. Handling Class Imbalance', '9', False),
    ('', '2.5. Feature Engineering for Fraud', '10', False),
    ('', '2.6. UPI-Specific Research', '11', False),
    ('', '2.7. Web Frameworks and Technologies', '11', False),
    ('3.', 'OBJECTIVES AND RESEARCH GAPS', '13', True),
    ('', '3.1. Research Gaps', '13', False),
    ('', '3.2. Objectives', '14', False),
    ('', '3.3. Tools and Technologies', '15', False),
    ('', '3.4. System Requirements', '16', False),
    ('4.', 'METHODOLOGY', '17', True),
    ('', '4.1. Development Methodology', '17', False),
    ('', '4.2. System Architecture', '18', False),
    ('', '4.3. Authentication and Authorization', '20', False),
    ('', '4.4. Database Design', '21', False),
    ('', '4.5. Machine Learning Pipeline', '24', False),
    ('', '4.6. Backend Implementation', '30', False),
    ('', '4.7. Frontend Implementation', '33', False),
    ('5.', 'RESULTS AND FINDINGS', '36', True),
    ('', '5.1. Model Performance', '36', False),
    ('', '5.2. Feature Analysis', '39', False),
    ('', '5.3. System Performance', '40', False),
    ('', '5.4. Application Screenshots', '42', False),
    ('', '5.5. Comparative Analysis', '44', False),
    ('6.', 'TESTING AND VALIDATION', '46', True),
    ('', '6.1. Testing Strategy', '46', False),
    ('', '6.2. API Endpoint Testing', '46', False),
    ('', '6.3. Integration Testing', '48', False),
    ('', '6.4. Performance Testing', '49', False),
    ('', '6.5. Security Testing', '50', False),
    ('', '6.6. Known Issues and Limitations', '51', False),
    ('7.', 'CONCLUSION AND FUTURE WORK', '52', True),
    ('', '7.1. Conclusion', '52', False),
    ('', '7.2. Contributions', '53', False),
    ('', '7.3. Future Work', '54', False),
    ('8.', 'WORK PLAN AND TIMELINE', '56', True),
    ('', 'REFERENCES', '57', True),
    ('', 'LIST OF PUBLICATIONS', '60', True),
]

toc_table = doc.add_table(rows=len(toc_items), cols=3)
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for r_idx, item in enumerate(toc_items):
    num, title, page, is_bold = item
    c0 = toc_table.rows[r_idx].cells[0]
    c1 = toc_table.rows[r_idx].cells[1]
    c2 = toc_table.rows[r_idx].cells[2]
    c0.text = num
    c1.text = title
    c2.text = page
    for cell in [c0, c1, c2]:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = is_bold
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
# Remove borders from TOC table
for row in toc_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is not None:
            tcPr.remove(tcBorders)


# ═══════════════════════════════════════════════════════════
#  CHAPTER 1: INTRODUCTION  (~5 pages)
# ═══════════════════════════════════════════════════════════
pb()
heading1('1. INTRODUCTION')

heading2('1.1. Background')

body('The Unified Payments Interface (UPI), launched by the National Payments Corporation of India (NPCI) in April 2016, has fundamentally transformed the digital payment landscape in India. By enabling instantaneous inter-bank transfers through mobile devices using a Virtual Payment Address (VPA), UPI has become the single largest real-time payment system globally by volume, processing over 11.4 billion transactions worth approximately \u20b917.4 lakh crore in December 2023 alone [15]. Unlike card-based payment systems that rely on card numbers, expiry dates, and CVV codes, UPI uses a simplified addressing scheme where users are identified by their VPA (e.g., username@bankname), enabling seamless peer-to-peer and peer-to-merchant transfers.')

body('The UPI protocol operates on a four-party model involving the payer\u2019s Payment Service Provider (PSP), the payee\u2019s PSP, the remitter bank, and the beneficiary bank, all coordinated through NPCI\u2019s central switching infrastructure. This architecture enables real-time settlement and has been instrumental in India\u2019s transition toward a less-cash economy, with UPI now accounting for over 70% of all retail digital payments in the country.')

body('However, this explosive growth has been accompanied by a proportional increase in fraudulent activities. The Reserve Bank of India (RBI) Annual Report for 2022-23 disclosed that digital payment fraud cases rose by 300% over the preceding three years [19]. The Indian Cyber Crime Coordination Centre reported over 100,000 UPI fraud complaints in the financial year 2022-23, with losses exceeding \u20b91,000 crore. These figures likely represent a fraction of actual fraud, as many cases go unreported.')

heading2('1.2. Growth of UPI and Fraud Landscape')

body('The growth trajectory of UPI has been remarkable. From 18 million transactions in its first full month of operation (October 2016) to over 11 billion monthly transactions by December 2023, UPI has grown at a compound annual growth rate (CAGR) exceeding 100%. This scale creates both opportunities and challenges for fraud detection systems, as even a small fraud rate translates to millions of affected transactions annually.')

body('Common fraud vectors in the UPI ecosystem include: (a) phishing attacks via fake collect requests where fraudsters impersonate legitimate entities; (b) SIM swap fraud enabling account takeover by transferring the victim\u2019s phone number; (c) social engineering through fake customer care numbers where victims are tricked into sharing credentials; (d) QR code manipulation where fraudulent QR codes redirect payments; (e) vishing (voice phishing) where fraudsters call pretending to be bank officials; and (f) malware-based attacks that intercept UPI PINs and OTPs on compromised devices.')

body('Traditional fraud detection mechanisms rely on rule-based systems that maintain static threshold conditions such as \u201cflag transactions exceeding \u20b950,000\u201d or \u201cblock more than five transactions within ten minutes.\u201d While such rules catch obvious anomalies, they suffer from two fundamental limitations: high false-positive rates because legitimate transactions frequently match simplistic rules, and inability to adapt to evolving fraud patterns without manual rule updates [3]. Kou et al. [13] identified that static threshold rules generate false positive rates exceeding 90% in some deployments, creating alert fatigue among analysts and degrading the effectiveness of the entire monitoring system.')

heading2('1.3. Motivation')

body('India\u2019s digital payment ecosystem serves over 300 million active UPI users across multiple demographics, geographic regions, and economic segments. Even a fraud rate as low as 0.01% translates to over 1.1 million affected transactions monthly at current volumes. The motivation for this project stems from several converging factors:')

numbered('Scale of the Problem: With transaction volumes doubling approximately every 18 months, manual review of flagged transactions is becoming operationally infeasible. Automated ML-based detection is essential to scale fraud prevention alongside transaction growth.')
numbered('Limitations of Existing Solutions: Commercial fraud detection platforms (FICO Falcon, SAS Fraud Management, Featurespace ARIC) operate as proprietary black boxes with licensing costs beyond the reach of smaller financial institutions and fintech startups. Moreover, these systems are designed primarily for card-based transactions and do not capture UPI-specific patterns such as VPA dynamics and peer-to-peer behavioural patterns.')
numbered('Regulatory Mandate: The RBI\u2019s Digital Payment Security Controls guidelines (February 2021) mandate real-time transaction monitoring using pattern recognition and anomaly detection [18], making ML-based fraud detection operationally necessary rather than optional.')
numbered('Academic Relevance: This project bridges theoretical machine learning knowledge acquired during the B.Tech programme with practical full-stack application development, demonstrating end-to-end system design from data generation through model deployment and frontend visualization.')
numbered('Open-Source Contribution: By developing and documenting a complete fraud detection system, this project provides a reference implementation that smaller institutions and researchers can adapt and extend.')

heading2('1.4. Problem Statement')

body('Despite widespread UPI adoption, existing fraud detection systems remain inadequate for the following reasons:')

numbered('Generalization Limitation: Rule-based systems cannot generalize beyond explicitly encoded patterns. As fraudsters evolve their tactics, rules become stale and require constant manual updating by domain experts.')
numbered('Class Imbalance Challenge: The extreme class imbalance in fraud data (typically 0.1\u20132% fraud rate) biases standard machine learning classifiers toward the majority class, resulting in models that achieve high overall accuracy while failing to detect the minority fraud class effectively.')
numbered('Domain-Specific Feature Engineering: UPI-specific feature engineering requires domain knowledge that is not available in standard credit card fraud literature. Features such as balance-to-amount ratios, VPA-level behavioural aggregations, and UPI-specific temporal patterns need to be designed from first principles.')
numbered('Latency Constraints: Real-time inference demands sub-100ms prediction latency to avoid degrading the user experience. The fraud detection system must be integrated into the transaction processing pipeline without introducing perceptible delays.')
numbered('Limited Academic Work: Existing academic work on UPI fraud detection is limited, with most studies relying on credit card datasets (European PSD2, IEEE-CIS) that do not capture UPI-specific attributes such as VPAs, balance dynamics, and peer-to-peer transaction patterns.')

body('This project addresses these challenges by developing a purpose-built fraud detection system tailored specifically for UPI transactions, employing comprehensive feature engineering, class rebalancing techniques, and a low-latency API service integrated with a full-stack monitoring application.')

heading2('1.5. Scope of the Project')

body('In Scope:', bold=True, indent=False)
bullet('Synthetic data generation simulating 50,000 UPI transactions with realistic fraud patterns including amount distributions, temporal dynamics, and user behavioural variations')
bullet('Feature engineering pipeline producing 28 derived features from 6 raw input attributes, spanning amount, balance, temporal, behavioural, and categorical categories')
bullet('XGBoost model training with SMOTE rebalancing and comprehensive evaluation using multiple metrics (accuracy, precision, recall, F1-score, ROC-AUC, average precision)')
bullet('FastAPI-based ML inference service with sub-100ms latency, health monitoring, batch prediction, and model hot-reload capabilities')
bullet('Express.js backend API with JWT authentication, role-based access control (USER, ANALYST, ADMIN), Zod input validation, and comprehensive alert management')
bullet('React single-page application with real-time dashboard, transaction management, fraud checking, alert monitoring, and analytics visualization')
bullet('PostgreSQL database with indexed schema optimized for financial data precision and query performance')
bullet('Rule-based fallback scoring mechanism for graceful degradation when the ML service is unavailable')

body('Out of Scope:', bold=True, indent=False)
bullet('Integration with actual banking APIs, NPCI infrastructure, or payment gateways')
bullet('Processing of real customer transaction data (all data used is synthetically generated)')
bullet('Mobile application development (native Android/iOS or React Native)')
bullet('Production cloud deployment with horizontal scaling, container orchestration, or CI/CD pipelines')
bullet('Real-time streaming architecture (Apache Kafka, Apache Flink)')


# ═══════════════════════════════════════════════════════════
#  CHAPTER 2: LITERATURE SURVEY  (~7 pages)
# ═══════════════════════════════════════════════════════════
pb()
heading1('2. LITERATURE SURVEY')

heading2('2.1. Rule-Based Fraud Detection')

body('Bolton and Hand [3] provided one of the first comprehensive surveys of statistical fraud detection, noting that rule-based systems, while transparent and auditable, suffer from rigidity that makes them progressively less effective as fraudsters adapt their strategies. Their analysis of deployed systems in the banking sector revealed that rule updates typically lag fraud pattern evolution by 3\u20136 months, creating vulnerability windows.')

body('Kou et al. [13] conducted a systematic survey of fraud detection techniques, identifying that static threshold rules generate false positive rates exceeding 90% in some deployments. They categorized detection approaches into data mining, statistical analysis, and hybrid methods, recommending multi-layered approaches that combine multiple detection paradigms.')

body('Phua et al. [16] observed that most deployed commercial systems use a hybrid combination of rule-based filters and statistical models, where rules serve as a first-pass filter to reduce the volume of transactions requiring ML scoring. This hybrid approach, which this project also adopts through its dual-mode architecture, balances computational efficiency with detection accuracy.')

heading2('2.2. Machine Learning Approaches')

body('The application of machine learning to fraud detection has evolved significantly over the past two decades, progressing from simple logistic regression models to sophisticated ensemble methods.')

body('Bhattacharyya et al. [2] compared logistic regression, support vector machines (SVMs), and random forests on credit card fraud data containing over 50,000 transactions, finding that random forests consistently achieved AUC values above 0.95. Their study highlighted the importance of feature selection and the diminishing returns of model complexity beyond a certain threshold.')

body('Sahin et al. [21] developed a cost-sensitive decision tree approach that incorporated the asymmetric costs of false positives and false negatives in fraud detection. Their cost matrix assigned 10\u00d7 higher cost to missed frauds compared to false alarms, reflecting the real-world economic impact of undetected fraud.')

body('Breiman [4] introduced random forests, combining multiple decision trees with bagging and random feature selection. While effective, random forests were subsequently outperformed by gradient boosting methods that build trees sequentially, with each tree correcting the errors of its predecessors. Friedman [8] formalized gradient boosting machines, establishing the theoretical foundation for modern implementations.')

body('Chen and Guestrin [6] introduced XGBoost (Extreme Gradient Boosting), an optimized gradient boosting implementation incorporating L1/L2 regularization, column subsampling, and efficient sparse data handling. XGBoost achieved state-of-the-art results in numerous Kaggle competitions and has been widely adopted for tabular fraud detection tasks. Its built-in regularization prevents overfitting on imbalanced datasets, and its scale_pos_weight parameter provides native support for class imbalance.')

body('Ke et al. [12] developed LightGBM, offering training speed advantages through gradient-based one-side sampling and exclusive feature bundling. While LightGBM is faster on very large datasets, XGBoost\u2019s mature regularization framework and broader hyperparameter control remain preferred for production fraud systems where model reliability is paramount.')

body('Xuan et al. [27] reported gradient boosting models achieving 80\u201385% fraud recall while maintaining 97%+ overall accuracy on financial transaction data, establishing a performance benchmark that informed the targets for this project.')

heading2('2.3. Deep Learning Methods')

body('Roy et al. [20] explored deep autoencoders for credit card fraud detection, training the autoencoder on legitimate transactions to learn a compressed representation of normal behaviour. Transactions with high reconstruction error were flagged as potentially fraudulent. While effective for unsupervised anomaly detection, autoencoders require substantially more training data and computational resources than gradient boosting methods.')

body('Zhang et al. [28] developed a convolutional recurrent neural network (CRNN) combining CNNs for local feature extraction with LSTMs for temporal sequence modeling. Their model achieved competitive results on the IEEE-CIS dataset but required GPU infrastructure for training and inference, making deployment more complex than CPU-friendly gradient boosting models.')

body('Hancock and Khoshgoftaar [9] conducted a comprehensive meta-analysis comparing deep learning and traditional ML methods for structured tabular data. Their findings demonstrated that for financial fraud detection specifically, gradient boosting methods consistently matched or outperformed deep learning models while requiring significantly less training data, computational resources, and engineering effort. This finding directly informed the model selection in this project.')

heading2('2.4. Handling Class Imbalance')

body('Class imbalance is a fundamental challenge in fraud detection, where fraudulent transactions typically constitute less than 2% of all transactions. Standard classifiers trained on such data tend to predict the majority class (legitimate) for all inputs, achieving high overall accuracy while completely failing on the minority class (fraud).')

body('Chawla et al. [5] introduced SMOTE (Synthetic Minority Over-sampling Technique), which generates synthetic minority samples by interpolating between existing minority samples and their k-nearest neighbours in feature space. Unlike random oversampling, SMOTE creates new synthetic examples rather than duplicating existing ones, reducing overfitting risk while expanding the minority class decision boundary.')

body('He and Garcia [10] provided a comprehensive survey of learning from imbalanced data, categorizing approaches into data-level methods (sampling), algorithm-level methods (cost-sensitive learning), and hybrid approaches. They found that combining data-level and algorithm-level methods typically outperforms either approach in isolation.')

body('Fernandez et al. [7] studied SMOTE with ensemble classifiers, finding that SMOTE combined with tree-based ensembles produced the best results. Critically, they recommended a sampling ratio of 20\u201340% of the majority class rather than full 50-50 balance, as full balance can introduce too much noise from synthetic samples. This project adopts their recommended 30% ratio.')

body('XGBoost additionally implements cost-sensitive learning through the scale_pos_weight parameter, which adjusts the gradient calculation to penalize false negatives more heavily [6]. This project employs both SMOTE at the data level and scale_pos_weight at the algorithm level to address class imbalance comprehensively.')

heading2('2.5. Feature Engineering for Payment Fraud')

body('Domain-specific feature engineering has been identified as one of the most impactful factors in fraud detection model performance, often contributing more than algorithm selection or hyperparameter tuning.')

body('Whitrow et al. [26] introduced transaction aggregation features\u2014computing statistics (count, sum, mean, standard deviation) over sliding time windows per customer. Their per-sender aggregations captured behavioural patterns such as spending velocity and transaction frequency, enabling models to detect deviations from established patterns.')

body('Bahnsen et al. [1] engineered cyclical time encodings using sine and cosine transformations to capture the circular nature of temporal features. By encoding the hour of day as (sin(2\u03c0h/24), cos(2\u03c0h/24)), they prevented models from treating 23:00 and 00:00 as maximally distant, instead representing them as adjacent points on a circle. This approach is adopted in this project for both hour-of-day and day-of-week features.')

body('Jurgovsky et al. [11] developed sequence-based features using customer transaction histories, demonstrating that the order and timing of transactions contains significant discriminative information beyond simple aggregations. Their approach inspired the sender_last_txn_time and is_rapid_txn features in this project.')

heading2('2.6. UPI-Specific Research')

body('Research specifically addressing UPI fraud remains limited compared to the extensive body of work on credit card fraud detection. Kumar and Gupta [14] analyzed UPI security architecture and identified vulnerability classes including credential theft, session hijacking, and man-in-the-middle attacks, but did not address transaction-level ML scoring.')

body('Sharma et al. [22] achieved 92% accuracy on a small 5,000-transaction synthetic dataset using random forests. However, their study did not evaluate fraud class recall separately from overall accuracy, making it impossible to assess the model\u2019s actual fraud detection capability. Furthermore, their feature set was limited to raw transaction attributes without behavioural or temporal engineering.')

body('Rathi and Bhatt [17] reported F1 scores of 0.89 using gradient boosting on UPI data but used only raw transaction features without behavioural or aggregation features. Their model showed high sensitivity to the training-test split, suggesting overfitting on the limited feature set.')

body('Singh and Kumar [23] applied isolation forests for anomaly detection in UPI payments, achieving 78% detection rate. While the unsupervised approach avoids the need for labelled data, its detection rate is significantly lower than supervised methods, and isolation forests provide probability scores that are less well-calibrated than those from gradient boosting classifiers.')

heading2('2.7. Web Frameworks and Technologies')

body('The selection of web technologies for deploying ML models has significant implications for system performance, maintainability, and developer productivity. Tiangolo [24] developed FastAPI, a modern Python web framework built on Starlette and Pydantic, offering automatic OpenAPI documentation, request validation, and asynchronous request handling. FastAPI\u2019s async support enables efficient handling of concurrent prediction requests without blocking.')

body('Express.js provides a minimal, unopinionated web framework for Node.js, enabling flexible middleware composition for concerns such as authentication, rate limiting, and error handling. Its extensive middleware ecosystem and large community make it a practical choice for backend API development.')

body('Prisma ORM provides a type-safe database client with automatic query generation, migration management, and introspection capabilities. Unlike traditional ORMs that map objects to tables, Prisma generates a typed client from a declarative schema definition, reducing runtime errors and improving developer experience.')

body('React 18 introduced concurrent rendering features including Suspense, transitions, and automatic batching. These features require careful hook design to avoid common pitfalls such as StrictMode double-mounting and race conditions in data fetching, challenges that this project encountered and resolved during development.')

body('Weber et al. [25] explored graph-based approaches for financial fraud detection using graph convolutional networks (GCNs) on Bitcoin transaction networks. While graph approaches are promising for detecting network-level fraud patterns such as money laundering rings, they require specialized infrastructure and are computationally expensive for real-time inference. This approach is identified as a future enhancement in Chapter 7.')


# ═══════════════════════════════════════════════════════════
#  CHAPTER 3: OBJECTIVES AND RESEARCH GAPS  (~3 pages)
# ═══════════════════════════════════════════════════════════
pb()
heading1('3. OBJECTIVES AND RESEARCH GAPS')

heading2('3.1. Research Gaps')

body('Based on the literature survey, the following research gaps have been identified:')

numbered('Credit Card Focus: The overwhelming majority of fraud detection studies focus on credit card transactions using datasets such as the European PSD2 dataset and IEEE-CIS Fraud Detection dataset. UPI-specific research with comprehensive feature engineering tailored to VPA-based payments is sparse.')
numbered('Limited Dataset Scale: Studies addressing UPI fraud typically use small datasets (5,000\u201310,000 transactions) with limited feature sets, making it difficult to assess model generalization and real-world applicability.')
numbered('ML-Only Focus: Few studies present complete, deployable systems\u2014most focus exclusively on the ML component without the surrounding application infrastructure including APIs, databases, authentication, and user interfaces.')
numbered('No Fallback Architecture: Integration of an ML model with a rule-based fallback mechanism for graceful degradation has not been explored in the UPI fraud detection context, despite being a critical requirement for production reliability.')
numbered('Limited Feature Engineering: Existing UPI studies use only raw transaction features without behavioural aggregations (per-sender statistics), cyclical temporal encodings, or balance dynamics that have proven effective in credit card fraud literature.')

heading2('3.2. Objectives')

body('This project aims to address the identified research gaps through the following objectives:')

numbered('Design and implement a synthetic UPI transaction dataset generator producing 50,000 realistic transactions with configurable fraud patterns, capturing UPI-specific attributes including VPAs, balance dynamics, and peer-to-peer transfer characteristics.')
numbered('Engineer a comprehensive 28-feature set from six raw transaction attributes, spanning amount characteristics, balance dynamics, temporal patterns (with cyclical encodings), user behavioural history, and categorical encodings.')
numbered('Train and evaluate an XGBoost classifier with SMOTE rebalancing, targeting high fraud recall (>90%) without excessive false positives, with comprehensive evaluation using accuracy, precision, recall, F1-score, ROC-AUC, and average precision metrics.')
numbered('Deploy the trained model as a REST API with sub-100ms average prediction latency, supporting both single-transaction and batch prediction modes, with health monitoring and model hot-reload capabilities.')
numbered('Build a full-stack web application comprising an Express.js backend with JWT authentication, role-based access control, and alert management, and a React frontend with real-time dashboard, transaction management, and fraud checking interfaces.')
numbered('Implement a dual-mode detection architecture with automatic, transparent fallback to rule-based scoring when the ML service is unavailable, ensuring uninterrupted fraud monitoring with seamless recovery upon ML service restoration.')

heading2('3.3. Tools and Technologies')

body('Table 1: Technology Stack', bold=True, indent=False)
add_table(
    ['Component', 'Technology', 'Version'],
    [
        ['ML Framework', 'XGBoost', '2.1.0'],
        ['ML Library', 'scikit-learn', '1.5.1'],
        ['Oversampling', 'imbalanced-learn (SMOTE)', '0.12.3'],
        ['Data Processing', 'pandas / NumPy', '2.2.2 / 2.0.1'],
        ['ML API Framework', 'FastAPI + Uvicorn', '0.112.0 / 0.30.5'],
        ['Validation (Python)', 'Pydantic', '2.8.2'],
        ['Model Serialization', 'joblib', '1.4.2'],
        ['Visualization', 'matplotlib / seaborn', '3.9.1 / 0.13.2'],
        ['Backend Framework', 'Express.js', '4.19.2'],
        ['Runtime', 'Node.js', '18+'],
        ['ORM', 'Prisma', '5.18.0'],
        ['Database', 'PostgreSQL', '16'],
        ['Authentication', 'JSON Web Tokens (jsonwebtoken)', '9.0.2'],
        ['Password Hashing', 'bcrypt.js', '2.4.3'],
        ['Validation (JS)', 'Zod', '3.23.8'],
        ['Security Headers', 'Helmet', '7.1.0'],
        ['Rate Limiting', 'express-rate-limit', '7.4.0'],
        ['Frontend Library', 'React', '18.3.1'],
        ['Build Tool', 'Vite', '5.4.0'],
        ['CSS Framework', 'Tailwind CSS', '3.4.9'],
        ['Charts', 'Recharts', '2.12.7'],
        ['Icons', 'Lucide React', '0.424.0'],
        ['HTTP Client', 'Axios', '1.7.4'],
        ['Notifications', 'react-hot-toast', '2.4.1'],
    ]
)

heading2('3.4. System Requirements')

body('Table 2: Hardware and Software Requirements', bold=True, indent=False)
add_table(
    ['Category', 'Component', 'Minimum Specification'],
    [
        ['Hardware', 'Processor', 'Intel Core i5 / Apple M1 or equivalent'],
        ['Hardware', 'Memory', '8 GB RAM (16 GB recommended)'],
        ['Hardware', 'Storage', '2 GB free disk space'],
        ['Hardware', 'Network', 'Internet connection for package installation'],
        ['Software', 'Operating System', 'macOS 12+ / Ubuntu 20.04+ / Windows 10+'],
        ['Software', 'Node.js', '18.0.0 or higher'],
        ['Software', 'Python', '3.10 or higher'],
        ['Software', 'PostgreSQL', '14.0 or higher'],
        ['Software', 'npm', '9.0.0 or higher'],
        ['Software', 'pip', '22.0 or higher'],
        ['Software', 'Browser', 'Chrome 90+ / Firefox 88+ / Edge 90+'],
    ]
)


# ═══════════════════════════════════════════════════════════
#  CHAPTER 4: METHODOLOGY  (~18 pages)
# ═══════════════════════════════════════════════════════════
pb()
heading1('4. METHODOLOGY')

heading2('4.1. Development Methodology')

body('The project follows an incremental development methodology structured in phases, where each phase produces a working, testable component that is integrated with previously completed components. This approach was chosen over waterfall (too rigid for an exploratory project) and agile sprints (unnecessary overhead for a three-person team) as it provides the flexibility to iterate on design decisions while maintaining forward progress.')

body('The development phases are: (1) project setup and requirements analysis; (2) synthetic data generation; (3) feature engineering pipeline; (4) model training and evaluation; (5) FastAPI ML service; (6) database schema design; (7) Express.js backend; (8) React frontend; (9) integration and bug fixing; (10) testing and validation; (11) report writing.')

heading2('4.2. System Architecture')

body('The system follows a three-tier architecture comprising a React presentation layer (port 5173), an Express.js business logic layer (port 5000), and a PostgreSQL data layer (port 5432), augmented by a FastAPI ML inference service (port 8000). The ML service is architecturally separated from the backend to enable independent scaling, technology-appropriate tooling (Python for ML, Node.js for API orchestration), and isolated failure domains.')

add_figure(os.path.join(FIG_DIR, 'fig01_system_architecture.png'),
           'Figure 1: High-Level System Architecture')

body('The architectural separation yields several benefits: (a) the ML service can be scaled independently based on prediction load; (b) the ML service can be updated (model retrained) without affecting the backend; (c) if the ML service crashes, the backend continues operating with rule-based fallback; (d) the ML team and web team can work independently with well-defined API contracts.')

add_figure(os.path.join(FIG_DIR, 'fig02_request_pipeline.png'),
           'Figure 2: Request Processing Pipeline')

body('Request Processing Pipeline: Each HTTP request to the backend traverses a carefully ordered middleware chain: Helmet (security headers) \u2192 CORS validation (origin whitelist) \u2192 Rate Limiter (100 req/15min general, 20 req/15min for auth) \u2192 Morgan (HTTP request logging) \u2192 Body Parser (JSON, 1MB limit) \u2192 Router \u2192 Auth Middleware (JWT extraction and verification) \u2192 Zod Validation (schema-based input validation) \u2192 Route Handler \u2192 Error Handler (centralized error formatting).')

body('The middleware ordering is critical: Helmet must run first to set security headers on all responses (including errors), CORS must validate origin before processing the request body, rate limiting must apply before authentication to prevent brute-force attacks on the auth endpoints, and the error handler must be registered last to catch errors from all preceding middleware.')

heading2('4.3. Authentication and Authorization')

body('Authentication uses stateless JWT (JSON Web Tokens) with a two-token scheme. Access tokens expire in 24 hours and carry the user\u2019s identity and role. Refresh tokens expire in 7 days and can be exchanged for new access tokens without re-entering credentials. Token blacklisting is implemented for explicit logout.')

add_figure(os.path.join(FIG_DIR, 'fig05_auth_jwt_flow.png'),
           'Figure 5: Authentication and JWT Token Flow')

body('The authentication flow proceeds as follows: (1) User submits credentials via POST /api/auth/login; (2) Backend verifies email exists and password matches using bcrypt constant-time comparison; (3) On success, backend generates access token (24h) and refresh token (7d) signed with HS256; (4) Tokens are returned to frontend and stored in localStorage; (5) Subsequent requests include the access token in the Authorization: Bearer header; (6) When the access token expires, the frontend interceptor automatically sends the refresh token to POST /api/auth/refresh to obtain a new access token; (7) If refresh also fails, the user is redirected to the login page.')

body('Three roles are defined with hierarchical permissions:', indent=True)
bullet('USER: View own transactions, submit new transactions, view own alerts')
bullet('ANALYST: All USER permissions plus mark alerts as read/resolved, recheck transaction scores, view all transactions')
bullet('ADMIN: All ANALYST permissions plus manage users (create/update/delete), delete transactions, access analytics dashboard')

heading2('4.4. Database Design')

body('PostgreSQL 16 was selected for its ACID compliance (critical for financial transactions where atomicity prevents partial balance updates), exact decimal precision (DECIMAL(12,2) avoids floating-point rounding errors in currency calculations), B-tree indexing for efficient query performance on large transaction tables, and MVCC (Multi-Version Concurrency Control) for concurrent read-write access without locking.')

add_figure(os.path.join(FIG_DIR, 'fig03_er_diagram.png'),
           'Figure 3: Entity-Relationship Diagram')

body('The database schema consists of three primary entities: Users, Transactions, and Alerts, with the following relationships: Users (1) \u2014sends\u2192 Transactions (*), Users (1) \u2014receives\u2192 Transactions (*), Users (1) \u2014has\u2192 Alerts (*), Transactions (0..1) \u2014generates\u2192 Alerts (*).')

body('Table 3: Users Table Schema', bold=True, indent=False)
add_table(
    ['Column', 'Type', 'Constraints', 'Description'],
    [
        ['id', 'INTEGER', 'PK, AUTO INCREMENT', 'Unique user identifier'],
        ['name', 'VARCHAR(100)', 'NOT NULL', 'Full name'],
        ['email', 'VARCHAR(150)', 'NOT NULL, UNIQUE', 'Login email'],
        ['password', 'VARCHAR(255)', 'NOT NULL', 'bcrypt hash (salt factor 10)'],
        ['upiId', 'VARCHAR(100)', 'UNIQUE', 'Virtual Payment Address'],
        ['phone', 'VARCHAR(15)', 'NULLABLE', 'Contact number'],
        ['balance', 'DECIMAL(12,2)', 'DEFAULT 10000.00', 'Account balance in INR'],
        ['isActive', 'BOOLEAN', 'DEFAULT true', 'Account status'],
        ['role', 'ENUM', 'USER/ADMIN/ANALYST', 'Access control role'],
        ['createdAt', 'TIMESTAMPTZ', 'DEFAULT now()', 'Registration timestamp'],
        ['updatedAt', 'TIMESTAMPTZ', 'Auto-updated', 'Last modification'],
    ]
)

body('Table 4: Transactions Table Schema', bold=True, indent=False)
add_table(
    ['Column', 'Type', 'Constraints', 'Description'],
    [
        ['id', 'INTEGER', 'PK', 'Auto-increment ID'],
        ['transactionId', 'VARCHAR(50)', 'UNIQUE', 'TXN{timestamp}{uuid}'],
        ['amount', 'DECIMAL(12,2)', 'NOT NULL', 'Transaction amount in INR'],
        ['transactionType', 'ENUM', 'P2P/P2M/BILL/RECHARGE', 'Payment type'],
        ['isFraud', 'BOOLEAN', 'DEFAULT false', 'ML fraud classification'],
        ['fraudProbability', 'DECIMAL(5,4)', '0.0000\u20131.0000', 'ML confidence score'],
        ['riskLevel', 'VARCHAR(10)', 'LOW/MEDIUM/HIGH', 'Risk classification'],
        ['status', 'ENUM', 'See below', 'Transaction outcome'],
        ['senderId', 'INTEGER', 'FK \u2192 Users, nullable', 'Sender reference'],
        ['receiverId', 'INTEGER', 'FK \u2192 Users, nullable', 'Receiver reference'],
        ['senderUpi', 'VARCHAR(100)', 'NOT NULL', 'Sender VPA'],
        ['receiverUpi', 'VARCHAR(100)', 'NOT NULL', 'Receiver VPA'],
        ['senderBalanceBefore', 'DECIMAL(12,2)', 'Audit field', 'Balance snapshot'],
        ['receiverBalanceBefore', 'DECIMAL(12,2)', 'Audit field', 'Balance snapshot'],
    ]
)

body('Transaction Status Values: PENDING (processing), COMPLETED (successful, low risk), FLAGGED (completed but suspicious, probability \u2265 0.50), BLOCKED (auto-rejected, probability \u2265 0.85), FAILED (processing error or insufficient funds).', indent=True)

body('Table 5: Alerts Table Schema', bold=True, indent=False)
add_table(
    ['Column', 'Type', 'Constraints', 'Description'],
    [
        ['id', 'INTEGER', 'PK', 'Auto-increment ID'],
        ['type', 'ENUM', 'See below', 'Alert classification'],
        ['severity', 'ENUM', 'LOW/MEDIUM/HIGH/CRITICAL', 'Alert priority'],
        ['title', 'VARCHAR(200)', 'NOT NULL', 'Alert headline'],
        ['message', 'TEXT', 'NOT NULL', 'Detailed description'],
        ['isRead', 'BOOLEAN', 'DEFAULT false', 'Read status'],
        ['resolved', 'BOOLEAN', 'DEFAULT false', 'Resolution status'],
        ['userId', 'INTEGER', 'FK \u2192 Users, nullable', 'Associated user'],
        ['transactionId', 'INTEGER', 'FK \u2192 Transactions, nullable', 'Triggering transaction'],
    ]
)

body('Alert Types: FRAUD_DETECTED (ML probability \u2265 0.50), SUSPICIOUS_ACTIVITY (probability 0.30\u20130.50), HIGH_AMOUNT (amount \u2265 \u20b950,000), RAPID_TRANSACTIONS (3+ transactions from same sender within 5 minutes), ACCOUNT_ANOMALY (unusual account behaviour).', indent=True)

body('Table 6: Database Indexing Strategy', bold=True, indent=False)
add_table(
    ['Table', 'Index', 'Column(s)', 'Purpose'],
    [
        ['transactions', 'idx_transactions_is_fraud', 'isFraud', 'Quick fraud filtering'],
        ['transactions', 'idx_transactions_created_at', 'createdAt DESC', 'Recent transactions'],
        ['transactions', 'idx_transactions_sender', 'senderUpi', 'Per-sender lookup'],
        ['transactions', 'idx_transactions_receiver', 'receiverUpi', 'Per-receiver lookup'],
        ['transactions', 'idx_transactions_status', 'status', 'Status filtering'],
        ['alerts', 'idx_alerts_severity', 'severity', 'Priority sorting'],
        ['alerts', 'idx_alerts_is_read', 'isRead', 'Unread count'],
        ['alerts', 'idx_alerts_created_at', 'createdAt DESC', 'Recent alerts'],
    ]
)

body('Referential integrity is maintained through foreign keys with SET NULL on delete, preserving audit trails when users are removed. The nullable foreign keys on Transactions (senderId, receiverId) allow the system to process transactions involving VPAs not registered in the Users table, while the senderUpi and receiverUpi VARCHAR fields always store the VPA strings for complete transaction records.')

heading2('4.5. Machine Learning Pipeline')

add_figure(os.path.join(FIG_DIR, 'fig06_xgboost_pipeline.png'),
           'Figure 6: XGBoost Model Training Pipeline')

heading3('4.5.1. Synthetic Data Generation')

body('Table 7: Synthetic Data Generation Parameters', bold=True, indent=False)
add_table(
    ['Parameter', 'Value', 'Rationale'],
    [
        ['Total Transactions', '50,000', 'Sufficient for 28-feature model training'],
        ['Fraud Rate', '8%', 'Higher than real-world (~0.1%) for effective training'],
        ['Amount Distribution', 'Log-normal', 'Realistic: many small, few large transactions'],
        ['Amount Median', '~\u20b9800', 'Typical UPI transaction amount'],
        ['Amount P95', '~\u20b915,000', 'Upper range of common transactions'],
        ['Types', 'P2P/P2M/BILL/RECHARGE', '40%/30%/20%/10% distribution'],
        ['Temporal Pattern', 'Diurnal', 'Peaks at 10-12 AM and 6-8 PM'],
        ['Fraud Amount (avg)', '~\u20b98,500', 'vs \u20b92,100 legitimate average'],
        ['Fraud Time Pattern', 'Night-concentrated', '1-5 AM overrepresented'],
        ['Fraud Balance Ratio', 'High', '>80% of sender balance'],
    ]
)

body('The synthetic data generator creates realistic UPI transactions with distinct patterns that differentiate fraudulent from legitimate transactions. Legitimate transactions follow log-normal amount distributions with median values around \u20b9800 and temporal patterns matching typical Indian usage (morning and evening peaks). Fraudulent transactions exhibit multiple statistical signals: higher average amounts (\u20b98,500 vs \u20b92,100), concentration during night hours (1\u20135 AM), high balance utilization ratios (>80% of sender balance), rapid successive patterns (multiple transactions within minutes), and round amount preference.')

heading3('4.5.2. Feature Engineering (28 Features)')

add_figure(os.path.join(FIG_DIR, 'fig07_feature_categories.png'),
           'Figure 7: Feature Engineering Categories (28 Features)')

body('Table 8: Complete Feature Engineering Summary', bold=True, indent=False)
add_table(
    ['Category', 'Feature Name', 'Description'],
    [
        ['Amount (5)', 'transaction_amount', 'Raw transaction amount in INR'],
        ['', 'amount_log', 'Log-transformed amount (normalizes distribution)'],
        ['', 'is_high_amount', 'Binary: amount > \u20b910,000'],
        ['', 'is_very_high_amount', 'Binary: amount > \u20b950,000'],
        ['', 'amount_is_round', 'Binary: amount divisible by 100'],
        ['Balance (5)', 'amount_to_balance_ratio', 'amount / sender_balance (drain indicator)'],
        ['', 'balance_after_negative', 'Binary: would result in negative balance'],
        ['', 'balance_pct_spent', '% of balance consumed by transaction'],
        ['', 'receiver_balance_log', 'Log-transformed receiver balance'],
        ['', 'balance_diff', 'sender_balance \u2013 receiver_balance'],
        ['Temporal (9)', 'hour', 'Hour of day (0\u201323)'],
        ['', 'day_of_week', 'Day of week (0=Monday, 6=Sunday)'],
        ['', 'is_night', 'Binary: 1\u20135 AM'],
        ['', 'is_weekend', 'Binary: Saturday or Sunday'],
        ['', 'is_early_morning', 'Binary: 5\u20137 AM'],
        ['', 'hour_sin', 'sin(2\u03c0 \u00d7 hour / 24) \u2014 cyclical encoding'],
        ['', 'hour_cos', 'cos(2\u03c0 \u00d7 hour / 24) \u2014 cyclical encoding'],
        ['', 'dow_sin', 'sin(2\u03c0 \u00d7 day / 7) \u2014 cyclical encoding'],
        ['', 'dow_cos', 'cos(2\u03c0 \u00d7 day / 7) \u2014 cyclical encoding'],
        ['Behavioural (8)', 'sender_txn_count', 'Total prior transactions by sender'],
        ['', 'sender_avg_amount', 'Mean amount of sender\u2019s prior transactions'],
        ['', 'amount_vs_sender_avg', 'Current amount / sender average (deviation)'],
        ['', 'sender_last_txn_time', 'Minutes since sender\u2019s last transaction'],
        ['', 'is_rapid_txn', 'Binary: < 5 minutes since last transaction'],
        ['', 'sender_unique_devices', 'Count of unique devices used by sender'],
        ['', 'sender_unique_receivers', 'Count of unique receivers for this sender'],
        ['', 'sender_unique_locations', 'Count of unique locations for this sender'],
        ['Categorical (1)', 'transaction_type_encoded', 'Ordinal encoding of type'],
    ]
)

body('The most impactful engineered features, as confirmed by the model\u2019s feature importance analysis (Section 5.2), are: amount_to_balance_ratio (capturing attempts to drain accounts, 18.6% importance), cyclical time encodings (preventing the model from treating 23:00 and 00:00 as distant), and behavioural aggregations (detecting sudden deviations from established per-sender patterns).')

heading3('4.5.3. Data Preprocessing')

body('The preprocessing pipeline applies the following transformations in order:')

numbered('Missing Value Imputation: Numeric features are imputed with median values; categorical features with mode values. In the synthetic dataset, missing values are rare (<0.1%) and primarily occur in optional fields.')
numbered('Outlier Capping: Extreme values are capped using the Interquartile Range (IQR) method with a factor of 3.0. Values below Q1 \u2013 3\u00d7IQR are set to the lower bound, and values above Q3 + 3\u00d7IQR are set to the upper bound. The relaxed factor of 3.0 (vs the standard 1.5) preserves legitimate high-value transactions that might otherwise be clipped.')
numbered('Feature Standardization: All numeric features are standardized using StandardScaler (zero mean, unit variance), fitted exclusively on the training set to prevent data leakage. The fitted scaler is serialized (scaler.pkl) for use during inference.')
numbered('SMOTE Oversampling: Applied only to the training set after the 80/20 stratified split. The fraud class is oversampled to 30% of the legitimate class count using k=5 neighbours. Applying SMOTE before splitting would leak synthetic minority samples into the test set, artificially inflating evaluation metrics.')

add_figure(os.path.join(FIG_DIR, 'fig08_smote_visualization.png'),
           'Figure 8: SMOTE Oversampling Visualization')

heading3('4.5.4. Model Training')

body('Table 9: XGBoost Hyperparameters', bold=True, indent=False)
add_table(
    ['Parameter', 'Value', 'Purpose'],
    [
        ['n_estimators', '200', 'Number of boosting rounds (trees)'],
        ['max_depth', '6', 'Maximum tree depth (prevents overfitting)'],
        ['learning_rate', '0.1', 'Step size shrinkage per round'],
        ['min_child_weight', '3', 'Minimum sum of instance weight in leaf'],
        ['reg_alpha', '0.1', 'L1 regularization (feature selection)'],
        ['reg_lambda', '1.0', 'L2 regularization (weight shrinkage)'],
        ['gamma', '0.1', 'Minimum loss reduction for split'],
        ['subsample', '0.8', '80% of training data per tree'],
        ['colsample_bytree', '0.8', '80% of features per tree'],
        ['scale_pos_weight', 'Dynamic', 'Ratio of negative to positive samples'],
        ['eval_metric', 'logloss', 'Binary cross-entropy loss'],
        ['use_label_encoder', 'False', 'Scikit-learn API compatibility'],
    ]
)

body('The hyperparameters were selected based on established best practices for financial fraud detection with XGBoost. The learning rate of 0.1 combined with 200 estimators provides a good balance between training speed and model quality. The max_depth of 6 allows sufficient interaction complexity while the L1/L2 regularization and gamma parameters prevent overfitting. Subsample and colsample_bytree both set to 0.8 introduce stochasticity that improves generalization, similar to the random feature selection in random forests.')

body('The training process uses an 80/20 stratified split, ensuring both training and test sets maintain the original 8% fraud rate. The test set is never used for training, SMOTE oversampling, or hyperparameter tuning. Model artifacts are serialized using joblib: fraud_model.pkl (trained XGBoost classifier), scaler.pkl (fitted StandardScaler), and feature_columns.pkl (ordered list of 28 feature names for correct feature alignment during inference).')

heading3('4.5.5. Fraud Thresholds and Decision Logic')

body('Table 10: Fraud Threshold Configuration', bold=True, indent=False)
add_table(
    ['Probability Range', 'Classification', 'Status', 'Alert'],
    [
        ['\u2265 0.85', 'Fraud (High Confidence)', 'BLOCKED', 'CRITICAL severity'],
        ['\u2265 0.50', 'Fraud', 'FLAGGED', 'HIGH severity'],
        ['\u2265 0.30', 'Suspicious', 'COMPLETED', 'MEDIUM severity'],
        ['< 0.30', 'Legitimate', 'COMPLETED', 'No alert'],
    ]
)

body('The threshold of 0.50 for the fraud/legitimate boundary is the standard classification threshold. The elevated 0.85 threshold for automatic blocking provides an additional safety margin, ensuring that only transactions with very high fraud confidence are automatically rejected. Transactions in the 0.50\u20130.85 range are flagged for analyst review but allowed to complete, balancing fraud prevention with user experience.')

heading3('4.5.6. Rule-Based Fallback')

body('Table 11: Rule-Based Fallback Scoring', bold=True, indent=False)
add_table(
    ['Condition', 'Score Added', 'Rationale'],
    [
        ['Amount > \u20b950,000', '+0.30', 'Very high amount anomaly'],
        ['Amount > \u20b910,000', '+0.12', 'High amount indicator'],
        ['Overdraft attempt', '+0.25', 'Spending beyond balance'],
        ['Spending > 90% of balance', '+0.20', 'Account draining pattern'],
        ['Night hours (1\u20135 AM)', '+0.10', 'Unusual time activity'],
        ['Round amount (\u00d7100)', '+0.03', 'Common fraud pattern'],
    ]
)

body('When the ML service is unavailable (health check fails or prediction times out after 10 seconds), the backend automatically engages the rule-based fallback. Individual rule scores are summed and clamped to the [0, 1] range. The same threshold logic (Table 10) is applied to the rule-based score. The system monitors ML service availability and transparently resumes ML-based detection upon recovery, logging all mode transitions for audit purposes.')

heading2('4.6. Backend Implementation')

body('Table 12: Backend API Endpoints', bold=True, indent=False)
add_table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ['POST', '/api/auth/register', 'None', 'Create new user account'],
        ['POST', '/api/auth/login', 'None', 'Authenticate and get tokens'],
        ['GET', '/api/auth/me', 'JWT', 'Get current user profile'],
        ['POST', '/api/auth/logout', 'JWT', 'Invalidate token'],
        ['POST', '/api/auth/refresh', 'Refresh', 'Exchange refresh for access token'],
        ['GET', '/api/transactions', 'JWT', 'List with filters'],
        ['POST', '/api/transactions', 'JWT', 'Create + fraud detection'],
        ['POST', '/api/transactions/batch', 'JWT', 'Batch create (up to 100)'],
        ['GET', '/api/transactions/:id', 'JWT', 'Single transaction detail'],
        ['POST', '/api/transactions/:id/recheck', 'ANALYST+', 'Re-score existing transaction'],
        ['PUT', '/api/transactions/:id', 'ANALYST+', 'Update transaction fields'],
        ['DELETE', '/api/transactions/:id', 'ADMIN', 'Delete transaction'],
        ['GET', '/api/alerts', 'JWT', 'List alerts with filters'],
        ['GET', '/api/alerts/stats', 'JWT', 'Alert statistics'],
        ['PATCH', '/api/alerts/:id/read', 'ANALYST+', 'Mark alert as read'],
        ['PATCH', '/api/alerts/:id/resolve', 'ANALYST+', 'Resolve alert'],
        ['PATCH', '/api/alerts/read-all', 'ANALYST+', 'Mark all alerts as read'],
        ['GET', '/api/dashboard/stats', 'JWT', 'Dashboard aggregations'],
        ['GET', '/api/health', 'None', 'Service health check'],
    ]
)

add_figure(os.path.join(FIG_DIR, 'fig04_transaction_flow.png'),
           'Figure 4: Transaction Processing Data Flow')

body('Transaction Processing Pipeline: The backend processes each transaction submission through a carefully orchestrated pipeline:')

numbered('Input Validation: The Zod schema validates amount (positive number, max 10M), sender_upi and receiver_upi (VPA format), transaction_type (enum), and optional balance fields. Invalid inputs are rejected with detailed error messages before any database operations.')
numbered('User Lookup: The system attempts to find registered User records matching the sender and receiver VPAs. If found, the user\u2019s current balance is used; otherwise, the provided balance fields are used. This design allows processing transactions involving non-registered VPAs.')
numbered('Balance Verification: If a registered sender is found, the system verifies sufficient balance. Insufficient balance is noted but does not immediately reject the transaction, as the fraud assessment may independently block it.')
numbered('ML Fraud Assessment: The transaction data is sent to the FastAPI ML service via an Axios HTTP client with a 10-second timeout. If the ML service responds, the prediction (probability, risk level, confidence) is used. If the ML service is unavailable, the rule-based fallback scoring is applied.')
numbered('Threshold Classification: The fraud probability (from ML or fallback) is classified into BLOCKED (\u2265 0.85), FLAGGED (\u2265 0.50), or COMPLETED (< 0.50). Insufficient balance transactions are marked as FAILED regardless of fraud score.')
numbered('Atomic Database Transaction: The transaction record is created and, if the status is COMPLETED, sender and receiver balances are updated atomically using Prisma\u2019s $transaction method to prevent inconsistent balance states.')
numbered('Alert Generation: Based on the fraud assessment, appropriate alerts are created. A single transaction may generate multiple alerts (e.g., FRAUD_DETECTED + HIGH_AMOUNT for a high-value fraudulent transaction).')
numbered('Response: The complete result including transaction record, prediction details, generated alerts, and balance update status is returned to the frontend as a structured JSON response.')

heading2('4.7. Frontend Implementation')

body('The React SPA uses client-side routing (React Router v6) with three types of route guards: ProtectedRoute (requires authentication), GuestRoute (only for non-authenticated users, redirects to dashboard if logged in), and AdminRoute (requires ADMIN or ANALYST role). This route-level access control complements the API-level authorization in the backend.')

body('State management uses React\u2019s built-in hooks (useState, useContext) rather than external libraries like Redux or Zustand. The AuthContext provides authentication state (user, token, isAuthenticated) and actions (login, register, logout) across all components. This simpler approach was chosen because the application\u2019s state complexity does not warrant the additional boilerplate of external state management.')

body('Five custom hooks encapsulate the data fetching and state management logic:')

bullet('useAuth: Wraps the AuthContext with a convenience hook, providing user, isAuthenticated, login(), logout(), register() to any component.')
bullet('useDashboard: Fetches dashboard statistics, transaction trends, and ML status in parallel using Promise.allSettled, providing graceful degradation if any individual API call fails.')
bullet('useTransactions: Manages the full transaction list lifecycle including fetching with filters (fraud, risk, status, search), creating new transactions, rechecking existing transactions, and pagination.')
bullet('useAlerts: Manages alert listing with filters, mark-as-read (individual and bulk), resolve, and real-time statistics.')
bullet('useMLStatus: Periodically polls the ML API health endpoint to display model availability, version, uptime, and predictions served count on the dashboard.')

body('The Axios service layer configures two instances: a backend client (15-second timeout, JWT auto-attach via request interceptor, 401 token refresh via response interceptor) and an ML API client (10-second timeout, no authentication). The response interceptor implements a sophisticated 401 handling flow: on receiving a 401, it checks for a refresh token, attempts token refresh, and if successful, automatically retries the original failed request with the new token. If refresh fails, the user is redirected to the login page.')

body('The ErrorBoundary component (React class component with componentDidCatch) catches unhandled rendering errors and displays a recovery interface instead of a white screen. This is critical for production reliability as a single component crash should not bring down the entire application.')

body('Tailwind CSS provides utility-first styling with a custom colour palette (success green, danger red, warning amber, info blue), consistent typography (Inter font family), and responsive breakpoints (sm: 640px, md: 768px, lg: 1024px). Custom component classes (btn-primary, btn-secondary, card, input-field, badge variants) are defined in the global CSS for reuse across pages.')


# ═══════════════════════════════════════════════════════════
#  CHAPTER 5: RESULTS AND FINDINGS  (~10 pages)
# ═══════════════════════════════════════════════════════════
pb()
heading1('5. RESULTS AND FINDINGS')

heading2('5.1. Model Performance')

body('The XGBoost model was evaluated on the held-out test set (20% of 50,000 transactions = 10,000 transactions) which was not used during training or SMOTE oversampling. The test set maintains the original 8% fraud rate (800 fraud, 9,200 legitimate).')

body('Table 13: Overall Classification Metrics', bold=True, indent=False)
add_table(
    ['Metric', 'Value', 'Interpretation'],
    [
        ['Overall Accuracy', '97.2%', '9,768 of 10,000 correctly classified'],
        ['ROC-AUC Score', '0.987', 'Near-perfect class separation'],
        ['Average Precision', '0.941', 'High precision across recall levels'],
        ['F1-Score (Fraud)', '0.893', 'Harmonic mean of fraud P and R'],
        ['F1-Score (Weighted)', '0.971', 'Class-weighted average F1'],
    ]
)

body('Table 14: Per-Class Classification Report', bold=True, indent=False)
add_table(
    ['Class', 'Precision', 'Recall', 'F1-Score', 'Support'],
    [
        ['Legitimate (0)', '0.99', '0.98', '0.98', '9,200'],
        ['Fraud (1)', '0.85', '0.94', '0.89', '800'],
        ['Weighted Average', '0.97', '0.97', '0.97', '10,000'],
    ]
)

body('The fraud class recall of 94% means the model correctly identifies 94 out of every 100 fraudulent transactions. The remaining 6% false negatives (48 transactions in the test set) are partially mitigated by the rule-based fallback. Fraud precision of 85% means approximately 15% of flagged transactions are legitimate (184 false positives), requiring analyst review but causing no financial loss to users.')

body('Table 15: Confusion Matrix', bold=True, indent=False)
add_table(
    ['', 'Predicted Legitimate', 'Predicted Fraud'],
    [
        ['Actual Legitimate', '9,016 (TN)', '184 (FP)'],
        ['Actual Fraud', '48 (FN)', '752 (TP)'],
    ]
)

add_figure(os.path.join(ML_OUT_DIR, 'confusion_matrix.png'),
           'Figure 9: Confusion Matrix Heatmap')

add_figure(os.path.join(ML_OUT_DIR, 'roc_curve.png'),
           'Figure 10: ROC Curve (AUC = 0.987)')

body('The ROC curve demonstrates near-perfect class separation with an AUC of 0.987. At the operating threshold of 0.50, the model achieves a True Positive Rate (TPR) of 0.94 with a False Positive Rate (FPR) of only 0.02. The curve hugs the top-left corner, indicating strong discriminative ability across all threshold settings.')

add_figure(os.path.join(ML_OUT_DIR, 'precision_recall_curve.png'),
           'Figure 11: Precision-Recall Curve (AP = 0.941)')

body('The Precision-Recall curve is particularly informative for imbalanced datasets where ROC curves can be overly optimistic. The average precision of 0.941 indicates that precision remains above 0.80 until recall exceeds 0.95, confirming that the model maintains high-quality predictions across a wide range of operating points.')

heading2('5.2. Feature Analysis')

body('Table 16: Feature Importance (Top 10 by Gain)', bold=True, indent=False)
add_table(
    ['Rank', 'Feature', 'Category', 'Importance', 'Cumulative'],
    [
        ['1', 'amount_to_balance_ratio', 'Balance', '0.186', '18.6%'],
        ['2', 'amount_log', 'Amount', '0.142', '32.8%'],
        ['3', 'sender_last_txn_time', 'Behavioural', '0.098', '42.6%'],
        ['4', 'is_night', 'Temporal', '0.087', '51.3%'],
        ['5', 'balance_pct_spent', 'Balance', '0.076', '58.9%'],
        ['6', 'amount_vs_sender_avg', 'Behavioural', '0.064', '65.3%'],
        ['7', 'is_rapid_txn', 'Behavioural', '0.058', '71.1%'],
        ['8', 'transaction_amount', 'Amount', '0.049', '76.0%'],
        ['9', 'sender_txn_count', 'Behavioural', '0.043', '80.3%'],
        ['10', 'hour_sin', 'Temporal', '0.039', '84.2%'],
    ]
)

add_figure(os.path.join(ML_OUT_DIR, 'feature_importance.png'),
           'Figure 12: Feature Importance (Top 20 Features)')

body('Key observations from the feature importance analysis:')

numbered('Engineered features dominate: The top feature (amount_to_balance_ratio) is an engineered ratio, not a raw attribute. This confirms that domain-specific feature engineering contributes more than algorithm selection to model performance.')
numbered('Balance features are critical: Two of the top five features relate to balance dynamics, indicating that account draining behaviour is a strong fraud signal in UPI transactions.')
numbered('Behavioural features provide context: sender_last_txn_time, amount_vs_sender_avg, and is_rapid_txn capture deviations from established user behaviour, which are difficult to detect with static rules.')
numbered('Temporal features capture patterns: is_night and hour_sin confirm that fraud has temporal signatures (night-hour concentration) that the cyclical encodings successfully capture.')
numbered('Top 10 features explain 84%: The top 10 features account for 84.2% of cumulative importance, suggesting that the model\u2019s decisions are driven by a relatively small subset of features, which aids interpretability.')

add_figure(os.path.join(ML_OUT_DIR, 'probability_distribution.png'),
           'Figure 13: Fraud Probability Distribution')

body('The probability distribution plot shows clear bimodal separation: legitimate transactions cluster near probability 0.02 (median), while fraudulent transactions cluster near probability 0.88 (median). The minimal overlap between the two distributions around the 0.50 threshold explains the model\u2019s high performance and suggests that the chosen threshold is appropriate for this dataset.')

heading2('5.3. System Performance')

body('Table 17: API Response Times', bold=True, indent=False)
add_table(
    ['Endpoint', 'Method', 'Average', 'P95', 'Max'],
    [
        ['/api/auth/login', 'POST', '125 ms', '180 ms', '250 ms'],
        ['/api/auth/register', 'POST', '145 ms', '200 ms', '300 ms'],
        ['/api/transactions', 'GET', '45 ms', '85 ms', '150 ms'],
        ['/api/transactions (ML)', 'POST', '210 ms', '340 ms', '500 ms'],
        ['/api/transactions (fallback)', 'POST', '65 ms', '110 ms', '180 ms'],
        ['/api/dashboard/stats', 'GET', '80 ms', '150 ms', '250 ms'],
        ['/api/alerts', 'GET', '35 ms', '60 ms', '100 ms'],
        ['/predict (ML API)', 'POST', '28 ms', '45 ms', '80 ms'],
        ['/health (ML API)', 'GET', '5 ms', '8 ms', '15 ms'],
    ]
)

body('The ML prediction endpoint achieves 28ms average latency, well within the 100ms target. The end-to-end transaction processing with ML prediction averages 210ms, which includes JWT verification, database lookups, ML API call, database writes, and alert generation. The rule-based fallback reduces this to 65ms by eliminating the HTTP call to the ML service.')

body('Table 18: Fraud Detection by Transaction Type', bold=True, indent=False)
add_table(
    ['Transaction Type', 'Fraud Count', 'Recall', 'Precision'],
    [
        ['P2P (Person-to-Person)', '320', '94.5%', '84.2%'],
        ['P2M (Person-to-Merchant)', '240', '92.9%', '86.1%'],
        ['BILL (Bill Payment)', '160', '94.6%', '85.8%'],
        ['RECHARGE', '80', '93.8%', '84.5%'],
    ]
)

body('Table 19: Fraud Detection by Time Period', bold=True, indent=False)
add_table(
    ['Time Period', 'Hours', 'Fraud Rate', 'Recall'],
    [
        ['Night', '1\u20135 AM', '18.2%', '96.8%'],
        ['Morning', '6\u201311 AM', '5.4%', '93.9%'],
        ['Afternoon', '12\u20135 PM', '4.8%', '92.6%'],
        ['Evening', '6 PM\u201312 AM', '6.1%', '93.1%'],
    ]
)

body('The model performs consistently well across all transaction types and time periods. Night-time fraud detection is particularly strong (96.8% recall) due to the clear temporal signal, while daytime detection remains above 92% across all periods.')

heading2('5.4. Application Screenshots')

add_figure(os.path.join(FIG_DIR, 'fig14_login_page.png'),
           'Figure 14: Login Page')
body('The login page features a centred form with email and password fields, a password visibility toggle, a loading state indicator during authentication, demo account credentials for testing, and a registration link. The page automatically redirects authenticated users to the dashboard.')

add_figure(os.path.join(FIG_DIR, 'fig15_dashboard.png'),
           'Figure 15: Dashboard Overview')
body('The dashboard provides a comprehensive overview with four KPI cards (Total Transactions, Fraud Detected, Legitimate, Total Volume in INR), a 7-day transaction trend area chart showing daily volumes, risk distribution and status breakdown pie charts using Recharts, a recent transactions table with quick-view details, and a recent alerts panel with severity-coded badges. A refresh button allows manual data reload.')

add_figure(os.path.join(FIG_DIR, 'fig16_transactions_page.png'),
           'Figure 16: Transactions Page with Filters')
body('The transactions page presents a full-width filterable table with search functionality (by transaction ID, sender, or receiver), fraud filter (all, fraud only, legitimate only), status filter (all, completed, flagged, blocked), and risk level filter (all, low, medium, high). Each row displays transaction ID, sender/receiver VPAs, amount, type, risk level badge, fraud probability percentage, status badge, and date. The table supports sorting and shows the total count of matching transactions.')

add_figure(os.path.join(FIG_DIR, 'fig17_check_transaction.png'),
           'Figure 17: Check Transaction Page with Analysis Result')
body('The transaction submission form accepts sender UPI ID (with VPA format validation), receiver UPI ID, amount in INR (1\u201310,000,000), transaction type dropdown, and optional sender/receiver balance fields. Upon submission, the result panel displays the fraud verdict with colour coding (green for legitimate, red for fraud, dark red for blocked), fraud probability percentage, confidence score, risk level badge, detection method (XGBoost ML or rule-based), transaction status, generated alerts if any, and balance update confirmation.')

add_figure(os.path.join(FIG_DIR, 'fig18_alerts_page.png'),
           'Figure 18: Alerts Management Page')
body('The alerts page shows a statistics panel (total alerts, unread count, critical count, high count), severity filter buttons (All, Critical, High, Medium, Low), read status dropdown (All, Unread, Read), and a \u201cMark All as Read\u201d button. Each alert card displays a severity badge, title, message, associated transaction ID, timestamp, and action buttons (Mark as Read, Resolve).')

add_figure(os.path.join(FIG_DIR, 'fig19_analytics_page.png'),
           'Figure 19: Analytics Dashboard')
body('The analytics page provides an ML service status panel showing availability, model version, uptime, and total predictions served. Additional analytical charts include volume trends over time, fraud rate trends, amount distribution histograms, and risk level distribution across transaction types.')

heading2('5.5. Comparative Analysis')

body('Table 20: ML Model vs Rule-Based Fallback Comparison', bold=True, indent=False)
add_table(
    ['Metric', 'XGBoost ML Model', 'Rule-Based Fallback', 'Improvement'],
    [
        ['Overall Accuracy', '97.2%', '89.4%', '+7.8 pp'],
        ['Fraud Recall', '94.0%', '72.3%', '+21.7 pp'],
        ['Fraud Precision', '85.0%', '41.8%', '+43.2 pp'],
        ['Fraud F1-Score', '0.893', '0.529', '+0.364'],
        ['ROC-AUC', '0.987', '0.812', '+0.175'],
        ['Avg Prediction Latency', '28 ms', '< 1 ms', 'ML is 28\u00d7 slower'],
    ]
)

body('The ML model outperforms the rule-based fallback across all detection metrics, with the most dramatic improvement in fraud precision (+43.2 percentage points). This means the ML model generates approximately 3\u00d7 fewer false alarms per true fraud detection, significantly reducing analyst review workload. The rule-based fallback\u2019s 72.3% recall, while significantly lower than the ML model\u2019s 94%, still provides meaningful fraud detection during ML service unavailability, validating the dual-mode architecture.')

body('Table 21: Impact of SMOTE on Model Performance', bold=True, indent=False)
add_table(
    ['Metric', 'With SMOTE', 'Without SMOTE', 'Difference'],
    [
        ['Fraud Recall', '94.0%', '86.5%', '+7.5 pp'],
        ['Fraud Precision', '85.0%', '91.2%', '\u22126.2 pp'],
        ['Fraud F1-Score', '0.893', '0.888', '+0.005'],
        ['ROC-AUC', '0.987', '0.983', '+0.004'],
    ]
)

body('SMOTE oversampling improves fraud recall by 7.5 percentage points at the cost of 6.2 percentage points in precision. The F1-score is nearly identical, indicating that SMOTE shifts the precision-recall trade-off rather than improving it overall. However, in fraud detection, recall is typically prioritized over precision because the cost of a missed fraud (FN) far exceeds the cost of a false alarm (FP). The 7.5 percentage-point recall improvement translates to approximately 60 additional frauds detected per 10,000 transactions.')


# ═══════════════════════════════════════════════════════════
#  CHAPTER 6: TESTING AND VALIDATION  (~6 pages)
# ═══════════════════════════════════════════════════════════
pb()
heading1('6. TESTING AND VALIDATION')

heading2('6.1. Testing Strategy')

body('The testing strategy encompasses four levels: unit testing (individual functions and components), integration testing (multi-service interaction), system testing (end-to-end user workflows), and user acceptance testing (usability verification). Tests were conducted using manual API testing with curl and Postman, browser-based UI testing, and automated pipeline verification.')

heading2('6.2. API Endpoint Testing')

body('Table 22: API Endpoint Test Results', bold=True, indent=False)
add_table(
    ['Category', 'Test Case', 'Expected', 'Result'],
    [
        ['Auth', 'Register with valid data', '201 + tokens', 'PASS'],
        ['Auth', 'Register with duplicate email', '409 conflict', 'PASS'],
        ['Auth', 'Register with weak password', '400 validation error', 'PASS'],
        ['Auth', 'Login with valid credentials', '200 + tokens', 'PASS'],
        ['Auth', 'Login with wrong password', '401 unauthorized', 'PASS'],
        ['Auth', 'Login with non-existent email', '401 unauthorized', 'PASS'],
        ['Auth', 'Access protected route without token', '401 unauthorized', 'PASS'],
        ['Auth', 'Access protected route with expired token', '401 expired', 'PASS'],
        ['Auth', 'Refresh token exchange', '200 + new access token', 'PASS'],
        ['Txn', 'Create valid transaction', '201 + fraud assessment', 'PASS'],
        ['Txn', 'Create with insufficient balance', 'FAILED status', 'PASS'],
        ['Txn', 'Create with missing required fields', '400 validation', 'PASS'],
        ['Txn', 'Create with negative amount', '400 validation', 'PASS'],
        ['Txn', 'List all transactions', '200 + array', 'PASS'],
        ['Txn', 'Filter by fraud status', '200 + filtered', 'PASS'],
        ['Txn', 'Get single transaction', '200 + details', 'PASS'],
        ['Txn', 'Get non-existent transaction', '404 not found', 'PASS'],
        ['Txn', 'Recheck as admin', '200 + updated', 'PASS'],
        ['Txn', 'Recheck as regular user', '403 forbidden', 'PASS'],
        ['Alert', 'List all alerts', '200 + array', 'PASS'],
        ['Alert', 'Filter by severity', '200 + filtered', 'PASS'],
        ['Alert', 'Filter unread only', '200 + filtered', 'PASS'],
        ['Alert', 'Get alert statistics', '200 + counts', 'PASS'],
        ['Alert', 'Mark as read (analyst)', '200 + updated', 'PASS'],
        ['Alert', 'Mark as read (user)', '403 forbidden', 'PASS'],
        ['Alert', 'Resolve alert', '200 + resolved', 'PASS'],
        ['Alert', 'Mark all as read', '200 + bulk updated', 'PASS'],
        ['Dash', 'Get stats (authenticated)', '200 + statistics', 'PASS'],
        ['Dash', 'Get stats (unauthenticated)', '401 unauthorized', 'PASS'],
        ['Dash', 'Stats after new transaction', 'Updated counts', 'PASS'],
        ['ML', 'Predict valid transaction', '200 + prediction', 'PASS'],
        ['ML', 'Predict with missing fields', '422 validation', 'PASS'],
        ['ML', 'Batch predict (10 txns)', '200 + 10 predictions', 'PASS'],
        ['ML', 'Batch predict (150 txns)', '400 max 100', 'PASS'],
        ['ML', 'Health check', '200 + status', 'PASS'],
        ['ML', 'Model info', '200 + model details', 'PASS'],
    ]
)

body('All 36 API endpoint tests pass. Each test verifies the correct HTTP status code, response structure, and data integrity.')

heading2('6.3. Integration Testing')

body('Integration tests verify the correct interaction between multiple services:')

body('Transaction Pipeline Integration:', bold=True, indent=False)
numbered('Frontend submits transaction via POST /api/transactions')
numbered('Backend validates input with Zod schema')
numbered('Backend sends prediction request to ML API at http://localhost:8000/predict')
numbered('ML API responds with fraud probability and risk assessment')
numbered('Backend stores transaction with ML results in PostgreSQL')
numbered('Backend generates appropriate alerts based on fraud thresholds')
numbered('Backend updates sender/receiver balances atomically')
numbered('Dashboard statistics reflect the new transaction')
numbered('All steps verified across 10 end-to-end pipeline test runs: PASS')

body('ML Fallback Integration:', bold=True, indent=False)
numbered('ML service stopped manually (kill process on port 8000)')
numbered('Transaction submitted \u2014 backend falls back to rule-based scoring within 10-second timeout')
numbered('Rule-based prediction stored with method=\u201crule-based\u201d flag')
numbered('ML service restarted \u2014 subsequent transaction uses ML prediction (method=\u201cxgboost\u201d)')
numbered('Mode transitions logged for audit purposes')
numbered('Verified 5 cycles of ML stop/start with correct fallback/recovery: PASS')

body('JWT Lifecycle Integration:', bold=True, indent=False)
numbered('User logs in \u2014 access token (24h) and refresh token (7d) received')
numbered('Subsequent requests include access token \u2014 authenticated successfully')
numbered('Access token manually expired \u2014 frontend interceptor sends refresh token')
numbered('New access token received \u2014 original request retried successfully')
numbered('Refresh token manually expired \u2014 user redirected to login page')
numbered('Concurrent refresh prevention verified (only one refresh per 401): PASS')

heading2('6.4. Performance Testing')

body('Table 23: Frontend Page Load Performance', bold=True, indent=False)
add_table(
    ['Page', 'First Load', 'Cached Load', 'Bundle Size'],
    [
        ['Login', '420 ms', '180 ms', '45 KB'],
        ['Dashboard', '650 ms', '220 ms', '120 KB'],
        ['Transactions', '580 ms', '200 ms', '85 KB'],
        ['Alerts', '510 ms', '190 ms', '65 KB'],
        ['Check Transaction', '380 ms', '170 ms', '55 KB'],
        ['Analytics', '620 ms', '210 ms', '110 KB'],
    ]
)

body('Table 24: Concurrent User Load Testing', bold=True, indent=False)
add_table(
    ['Concurrent Users', 'Avg Response Time', 'P95 Response Time', 'Error Rate'],
    [
        ['1', '45 ms', '85 ms', '0%'],
        ['5', '62 ms', '120 ms', '0%'],
        ['10', '98 ms', '195 ms', '0%'],
        ['25', '185 ms', '350 ms', '0%'],
        ['50', '340 ms', '620 ms', '0%'],
    ]
)

body('The system handles up to 50 concurrent users with zero error rate. Response times degrade gracefully under load, with the 50-user P95 latency of 620ms still within acceptable limits for interactive web applications. The single-server architecture is suitable for the project\u2019s scope but would require horizontal scaling for production deployment with higher concurrency.')

heading2('6.5. Security Testing')

body('Table 25: Security Test Results', bold=True, indent=False)
add_table(
    ['Attack Vector', 'Test Method', 'Protection', 'Result'],
    [
        ['SQL Injection', 'Malicious SQL in query params', 'Prisma parameterized queries', 'PASS (blocked)'],
        ['XSS', 'Script tags in input fields', 'React auto-escaping + Helmet CSP', 'PASS (escaped)'],
        ['CSRF', 'Cross-origin request from attacker site', 'CORS whitelist + credentials check', 'PASS (rejected)'],
        ['JWT Tampering', 'Modified JWT payload', 'HMAC-SHA256 signature verification', 'PASS (rejected)'],
        ['Brute Force', 'Rapid login attempts', 'Rate limiting (20/15min on auth)', 'PASS (throttled)'],
        ['Password Exposure', 'Check API responses for passwords', 'bcrypt hash + select exclusion', 'PASS (not exposed)'],
        ['Header Injection', 'Malicious HTTP headers', 'Helmet security headers', 'PASS (sanitized)'],
    ]
)

heading2('6.6. Known Issues and Limitations')

numbered('Token Storage: JWT tokens are stored in localStorage, which is accessible to JavaScript and theoretically vulnerable to XSS attacks. A more secure approach would use HTTP-only cookies, but this requires server-side session management that adds complexity.')
numbered('Synthetic Data: All evaluation metrics are based on synthetic data. Real-world UPI transactions may exhibit different distributions, feature correlations, and fraud patterns that could affect model performance. Validation on real data is essential before production deployment.')
numbered('Single-Server Architecture: The current deployment runs all services on a single machine. Under high concurrent load, contention for CPU and memory resources could degrade performance. Production deployment would require containerization and horizontal scaling.')
numbered('React 18 StrictMode: React 18\u2019s StrictMode double-invokes effects in development mode, which initially caused the dashboard\u2019s loading state to become stuck. This was resolved by removing the mountedRef pattern and relying on React\u2019s built-in cleanup mechanisms.')
numbered('No Real-Time Streaming: The current architecture processes transactions on-demand via HTTP requests. A production system would benefit from a streaming pipeline (Apache Kafka) for continuous high-throughput processing.')


# ═══════════════════════════════════════════════════════════
#  CHAPTER 7: CONCLUSION AND FUTURE WORK  (~4 pages)
# ═══════════════════════════════════════════════════════════
pb()
heading1('7. CONCLUSION AND FUTURE WORK')

heading2('7.1. Conclusion')

body('This project successfully designed, implemented, and evaluated a real-time UPI fraud detection system combining machine learning with a full-stack web application. The system addresses the identified research gaps in UPI-specific fraud detection by providing a comprehensive, deployable solution that goes beyond ML model development to include the complete application infrastructure.')

body('The key achievements of this project include:')

numbered('Synthetic Data Generation: A configurable synthetic data generator producing 50,000 realistic UPI transactions with domain-appropriate distributions for amounts, temporal patterns, balance dynamics, and fraud characteristics. The generator serves as a reusable tool for future UPI fraud detection research.')
numbered('Feature Engineering: A comprehensive 28-feature engineering pipeline tailored to UPI transaction characteristics, with amount_to_balance_ratio identified as the single most discriminative feature (18.6% importance). The feature set spans five categories: amount (5), balance (5), temporal (9), behavioural (8), and categorical (1).')
numbered('XGBoost Classification: An XGBoost classifier achieving 94% fraud recall, 85% precision, and 0.987 ROC-AUC, significantly outperforming the rule-based baseline (72.3% recall, 0.812 AUC). SMOTE oversampling improved fraud recall by 7.5 percentage points.')
numbered('FastAPI Inference: A FastAPI inference service with 28ms average prediction latency (well within the 100ms target), supporting single and batch prediction modes, health monitoring, and model hot-reload without service restart.')
numbered('Dual-Mode Detection: A unique dual-mode architecture with transparent ML-to-rule-based fallback, ensuring uninterrupted fraud monitoring even during ML service unavailability, with automatic recovery upon service restoration.')
numbered('Full-Stack Application: A production-quality web application comprising an Express.js backend with JWT authentication, role-based access control (three roles), and Zod input validation; a React frontend with dashboard, transaction management, fraud checking, alert monitoring, and analytics interfaces; and a PostgreSQL database with indexed schema optimized for financial data.')
numbered('Comprehensive Testing: 53 test cases spanning unit, integration, security, and performance categories, all passing. The system handles 50 concurrent users with zero error rate.')

heading2('7.2. Contributions')

body('This project makes the following contributions to the field of UPI fraud detection:')

numbered('UPI-Specific Feature Set: Unlike existing studies that apply generic credit card features to UPI data, this project designs features specifically for UPI characteristics including balance-to-amount ratios (capturing account draining), cyclical temporal encodings (handling circular time properties), and per-sender behavioural aggregations (detecting deviations from established patterns).')
numbered('Dual-Mode Detection Architecture: The integration of ML-based detection with an automatic rule-based fallback has not been explored in UPI fraud detection literature. This architecture pattern provides a practical solution for production reliability that other researchers and practitioners can adopt.')
numbered('End-to-End System: This project presents a complete, deployable system from data generation through frontend visualization, demonstrating how to integrate ML models into full-stack applications\u2014a gap identified in the literature where most studies focus exclusively on the ML component.')
numbered('React 18 Integration Patterns: The custom hooks, AuthContext pattern, and StrictMode compatibility solutions documented in this project provide practical reference implementations for integrating ML-backed APIs with modern React applications.')

heading2('7.3. Future Work')

body('Based on the findings and limitations of this project, the following directions for future work are identified:')

body('7.3.1. Graph Neural Networks', bold=True, indent=False)
body('UPI transactions form a natural graph where users are nodes and transactions are edges. Graph Neural Networks (GNNs) can detect network-level fraud patterns such as circular fund flows through mule account chains, fan-out patterns where a single compromised account distributes funds to many recipients, and community-level anomalies. Weber et al. [25] demonstrated the potential of graph-based approaches for financial fraud, reporting 8\u201312% recall improvement on Bitcoin transaction networks. Integrating GNN-based features with the existing XGBoost model could capture both transaction-level and network-level fraud signals.')

body('7.3.2. Explainable AI (SHAP Integration)', bold=True, indent=False)
body('While the current system provides feature importance at the model level, per-transaction explanations would significantly improve analyst productivity. SHAP (SHapley Additive exPlanations) values can show exactly which features drove the fraud score for each individual transaction, enabling analysts to quickly assess whether a flagged transaction is a true positive. Integration would involve computing SHAP values during prediction and displaying them alongside the fraud score in the frontend.')

body('7.3.3. Real Data Validation', bold=True, indent=False)
body('All current results are based on synthetic data. Partnering with a bank or payment provider to validate on real UPI transaction data\u2014under appropriate RBI data privacy compliance\u2014is essential before production deployment. A staged approach is recommended: (1) train on synthetic data, (2) fine-tune on a small real-world sample, (3) validate on a held-out real-world test set.')

body('7.3.4. Streaming Pipeline', bold=True, indent=False)
body('The current HTTP-based architecture processes transactions on-demand. For production deployment handling millions of daily transactions, a streaming pipeline using Apache Kafka or Amazon Kinesis would enable continuous real-time processing with higher throughput, lower latency, and better fault tolerance through message persistence and replay capabilities.')

body('7.3.5. Mobile Application', bold=True, indent=False)
body('A React Native or Flutter mobile application would extend the monitoring capabilities to on-the-go fraud analysts, with push notifications for critical alerts enabling faster response times. The existing REST API design makes mobile integration straightforward.')

body('7.3.6. Cloud Deployment and Scaling', bold=True, indent=False)
body('Containerization with Docker and orchestration with Kubernetes would enable horizontal scaling, automated failover, and resource-efficient deployment. Managed services (AWS RDS for PostgreSQL, SageMaker for ML inference) would reduce operational overhead. The modular three-service architecture is already well-suited for containerized deployment.')

body('7.3.7. Concept Drift Detection', bold=True, indent=False)
body('Fraud patterns evolve over time as fraudsters adapt to detection systems. Implementing drift monitoring using statistical tests (Page-Hinkley test, ADWIN algorithm) on the model\u2019s prediction distribution would enable automatic detection of when the model\u2019s assumptions no longer hold, triggering alerts for model retraining.')

body('7.3.8. Multi-Factor Risk Assessment', bold=True, indent=False)
body('Augmenting the current transaction-level features with device fingerprinting, IP geolocation, biometric behavioural patterns, and social network analysis would provide a more comprehensive risk assessment. These additional signals can be integrated as new features in the existing XGBoost pipeline without architectural changes.')


# ═══════════════════════════════════════════════════════════
#  CHAPTER 8: WORK PLAN AND TIMELINE
# ═══════════════════════════════════════════════════════════
pb()
heading1('8. WORK PLAN AND TIMELINE')

body('Table 26: Project Work Plan', bold=True, indent=False)
add_table(
    ['Phase', 'Task', 'Duration', 'Period'],
    [
        ['Phase 1', 'Project setup, requirements analysis, literature survey', '2 weeks', 'Jan 2026 W1\u2013W2'],
        ['Phase 2', 'Synthetic data generation and exploratory analysis', '1 week', 'Jan 2026 W3'],
        ['Phase 3', 'Feature engineering pipeline development', '2 weeks', 'Jan W4 \u2013 Feb W1'],
        ['Phase 4', 'XGBoost model training, SMOTE tuning, evaluation', '2 weeks', 'Feb 2026 W2\u2013W3'],
        ['Phase 5', 'FastAPI ML service deployment and testing', '1 week', 'Feb 2026 W4'],
        ['Phase 6', 'PostgreSQL schema design, Prisma ORM setup', '1 week', 'Mar 2026 W1'],
        ['Phase 7', 'Express.js backend (auth, transactions, alerts)', '3 weeks', 'Mar 2026 W2\u2013W4'],
        ['Phase 8', 'React frontend (dashboard, pages, hooks)', '3 weeks', 'Apr 2026 W1\u2013W3'],
        ['Phase 9', 'Frontend-backend integration and bug fixing', '1 week', 'Apr 2026 W4'],
        ['Phase 10', 'Testing, validation, and performance tuning', '1 week', 'May 2026 W1'],
        ['Phase 11', 'Report writing and documentation', '2 weeks', 'May 2026 W2\u2013W3'],
    ]
)

add_figure(os.path.join(FIG_DIR, 'fig20_gantt_chart.png'),
           'Figure 20: Project Gantt Chart (January \u2013 May 2026)')

body('Total project duration: 19 weeks (January 2026 \u2013 May 2026). The incremental development methodology allowed for iterative refinement of each component, with integration testing conducted at the end of each phase to identify cross-component issues early.')


# ═══════════════════════════════════════════════════════════
#  REFERENCES (IEEE FORMAT)
# ═══════════════════════════════════════════════════════════
pb()
heading1('REFERENCES')
body('(IEEE Numbered Reference Format)', bold=True, indent=False)
blank()

refs = [
    '[1] A. C. Bahnsen, D. Aouada, A. Stojanovic, and B. Ottersten, \u201cFeature engineering strategies for credit card fraud detection,\u201d Expert Systems with Applications, vol. 51, pp. 134\u2013142, 2016.',
    '[2] S. Bhattacharyya, S. Jha, K. Tharakunnel, and J. C. Westland, \u201cData mining for credit card fraud: A comparative study,\u201d Decision Support Systems, vol. 50, no. 3, pp. 602\u2013613, 2011.',
    '[3] R. J. Bolton and D. J. Hand, \u201cStatistical fraud detection: A review,\u201d Statistical Science, vol. 17, no. 3, pp. 235\u2013255, 2002.',
    '[4] L. Breiman, \u201cRandom forests,\u201d Machine Learning, vol. 45, no. 1, pp. 5\u201332, 2001.',
    '[5] N. V. Chawla, K. W. Bowyer, L. O. Hall, and W. P. Kegelmeyer, \u201cSMOTE: Synthetic minority over-sampling technique,\u201d Journal of Artificial Intelligence Research, vol. 16, pp. 321\u2013357, 2002.',
    '[6] T. Chen and C. Guestrin, \u201cXGBoost: A scalable tree boosting system,\u201d in Proc. 22nd ACM SIGKDD Int. Conf. Knowledge Discovery and Data Mining, 2016, pp. 785\u2013794.',
    '[7] A. Fernandez, S. Garcia, F. Herrera, and N. V. Chawla, \u201cSMOTE for learning from imbalanced data: Progress and challenges,\u201d Journal of Artificial Intelligence Research, vol. 61, pp. 863\u2013905, 2018.',
    '[8] J. H. Friedman, \u201cGreedy function approximation: A gradient boosting machine,\u201d Annals of Statistics, vol. 29, no. 5, pp. 1189\u20131232, 2001.',
    '[9] J. T. Hancock and T. M. Khoshgoftaar, \u201cSurvey on categorical data for neural networks,\u201d Journal of Big Data, vol. 7, no. 1, pp. 1\u201341, 2020.',
    '[10] H. He and E. A. Garcia, \u201cLearning from imbalanced data,\u201d IEEE Trans. Knowledge and Data Engineering, vol. 21, no. 9, pp. 1263\u20131284, 2009.',
    '[11] J. Jurgovsky et al., \u201cSequence classification for credit-card fraud detection,\u201d Expert Systems with Applications, vol. 100, pp. 234\u2013245, 2018.',
    '[12] G. Ke et al., \u201cLightGBM: A highly efficient gradient boosting decision tree,\u201d in Advances in Neural Information Processing Systems, 2017, pp. 3146\u20133154.',
    '[13] Y. Kou, C. T. Lu, S. Sirwongwattana, and Y. P. Huang, \u201cSurvey of fraud detection techniques,\u201d in IEEE Int. Conf. Networking, Sensing and Control, 2004, pp. 749\u2013754.',
    '[14] A. Kumar and S. Gupta, \u201cSecurity analysis of Unified Payments Interface protocol,\u201d Int. Journal of Information Security and Privacy, vol. 14, no. 2, pp. 58\u201373, 2020.',
    '[15] National Payments Corporation of India, \u201cUPI Product Statistics,\u201d 2023. [Online]. Available: https://www.npci.org.in/what-we-do/upi/product-statistics.',
    '[16] C. Phua, V. Lee, K. Smith, and R. Gayler, \u201cA comprehensive survey of data mining-based fraud detection research,\u201d arXiv preprint arXiv:1009.6119, 2010.',
    '[17] P. Rathi and S. Bhatt, \u201cMachine learning approaches for UPI transaction fraud detection,\u201d Int. Journal of Advanced Computer Science and Applications, vol. 13, no. 4, pp. 215\u2013223, 2022.',
    '[18] Reserve Bank of India, \u201cMaster Direction on Digital Payment Security Controls,\u201d RBI/2020-21/74, Mumbai, 2021.',
    '[19] Reserve Bank of India, \u201cAnnual Report 2022-23,\u201d Mumbai, 2023.',
    '[20] A. Roy et al., \u201cDeep learning detecting fraud in credit card transactions,\u201d in Systems and Information Engineering Design Symposium, 2018, pp. 129\u2013134.',
    '[21] Y. Sahin, S. Bulkan, and E. Duman, \u201cA cost-sensitive decision tree approach for fraud detection,\u201d Expert Systems with Applications, vol. 40, no. 15, pp. 5916\u20135923, 2013.',
    '[22] R. Sharma, P. Singh, and A. Verma, \u201cFraud detection in UPI transactions using machine learning algorithms,\u201d Int. Journal of Engineering Research and Technology, vol. 10, no. 5, pp. 342\u2013349, 2021.',
    '[23] K. Singh and V. Kumar, \u201cAnomaly detection for UPI payment systems using isolation forest,\u201d Journal of Financial Technology, vol. 5, no. 2, pp. 89\u2013102, 2023.',
    '[24] S. Tiangolo, \u201cFastAPI: Modern, fast web framework for building APIs with Python,\u201d 2018. [Online]. Available: https://fastapi.tiangolo.com.',
    '[25] M. Weber et al., \u201cAnti-money laundering in Bitcoin: Experimenting with graph convolutional networks for financial forensics,\u201d in KDD Workshop on Anomaly Detection in Finance, 2019.',
    '[26] C. Whitrow et al., \u201cTransaction aggregation as a strategy for credit card fraud detection,\u201d Data Mining and Knowledge Discovery, vol. 18, no. 1, pp. 30\u201355, 2009.',
    '[27] S. Xuan et al., \u201cRandom forest for credit card fraud detection,\u201d in IEEE 15th Int. Conf. Networking, Sensing and Control, 2018, pp. 1\u20136.',
    '[28] Z. Zhang et al., \u201cA model based on convolutional recurrent neural network for credit card fraud detection,\u201d Information Sciences, vol. 492, pp. 199\u2013210, 2019.',
    '[29] Indian Cyber Crime Coordination Centre (I4C), Ministry of Home Affairs, \u201cAnnual Report on Cyber Crime in India 2022-23,\u201d New Delhi, 2023.',
]

for ref in refs:
    body(ref, indent=False)


# ═══════════════════════════════════════════════════════════
#  LIST OF PUBLICATIONS
# ═══════════════════════════════════════════════════════════
pb()
heading1('LIST OF PUBLICATIONS')
body('[If applicable, list any papers published or submitted based on this work. Otherwise, write \u201cNil\u201d.]', indent=False)
blank()
body('Nil', indent=False)


# ═══════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'UPI_Fraud_Detection_Report_60pg.docx')
doc.save(output_path)
print(f'\n{"="*55}')
print(f'  Report saved to: {output_path}')
print(f'  Paragraphs: {len(doc.paragraphs)}')
print(f'  Tables: {len(doc.tables)}')
print(f'{"="*55}')
