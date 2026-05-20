#!/usr/bin/env python3
"""
Generate a proper Word (.docx) report from markdown chapter files.
Produces a formatted DTU B.Tech Major Project Report.
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# ── Page Setup ────────────────────────────────────────────
for section in doc.sections:
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.81)  # 1.5 inches
    section.right_margin = Cm(2.54)

# ── Style Setup ───────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# Helper functions
def add_page_break():
    doc.add_page_break()

def add_centered(text, size=16, bold=True, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = 'Times New Roman'
    return p

def add_body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_table_from_data(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
    return table

def add_empty_lines(n=1):
    for _ in range(n):
        doc.add_paragraph()


# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════
add_empty_lines(3)
add_centered('DELHI TECHNOLOGICAL UNIVERSITY', 18, True, 0, 4)
add_centered('(Formerly Delhi College of Engineering)', 14, False, 0, 4)
add_centered('Bawana Road, Delhi — 110042', 12, False, 0, 24)

add_centered('MAJOR PROJECT REPORT', 18, True, 24, 24)

add_centered('Real-Time UPI Fraud Detection System\nUsing Machine Learning', 16, True, 12, 24)

add_centered('Submitted in partial fulfilment of the requirements\nfor the award of the degree of', 12, False, 12, 12)
add_centered('Bachelor of Technology', 14, True, 6, 4)
add_centered('in', 12, False, 0, 4)
add_centered('[Department Name]', 14, True, 0, 24)

add_centered('Submitted by:', 12, True, 12, 6)
add_centered('[Student Name 1] — [Roll Number]\n[Student Name 2] — [Roll Number]\n[Student Name 3] — [Roll Number]\n[Student Name 4] — [Roll Number]', 12, False, 0, 18)

add_centered('Under the supervision of:', 12, True, 12, 6)
add_centered('[Supervisor Name]\n[Designation]\nDepartment of [Department Name]\nDelhi Technological University', 12, False, 0, 24)

add_centered('Session: 2024-25', 12, True, 12, 0)

# ═══════════════════════════════════════════════════════════
# CERTIFICATE
# ═══════════════════════════════════════════════════════════
add_page_break()
add_centered('CERTIFICATE', 16, True, 24, 18)

add_body('This is to certify that the Major Project Report entitled "Real-Time UPI Fraud Detection System Using Machine Learning" submitted by [Student Name 1] (Roll No: [Roll Number]), [Student Name 2] (Roll No: [Roll Number]), [Student Name 3] (Roll No: [Roll Number]), and [Student Name 4] (Roll No: [Roll Number]) in partial fulfilment of the requirements for the award of the degree of Bachelor of Technology in [Department Name] from Delhi Technological University, Delhi, is an authentic record of work carried out by them under my supervision and guidance.', indent=True)

add_body('The matter embodied in this report has not been submitted for the award of any other degree or diploma.', indent=True)

add_empty_lines(6)
add_body('Date:')
add_empty_lines(2)
add_body('[Supervisor Name]', bold=True)
add_body('[Designation]')
add_body('Department of [Department Name]')
add_body('Delhi Technological University, Delhi')

# ═══════════════════════════════════════════════════════════
# DECLARATION
# ═══════════════════════════════════════════════════════════
add_page_break()
add_centered('DECLARATION', 16, True, 24, 18)

add_body('We hereby declare that the Major Project Report entitled "Real-Time UPI Fraud Detection System Using Machine Learning" submitted by us to the Department of [Department Name], Delhi Technological University, Delhi, in partial fulfilment of the requirements for the award of the degree of Bachelor of Technology, is a bona fide record of original work carried out by us. The matter embodied in this report has not been submitted for the award of any other degree or diploma of any university or institution.', indent=True)

add_empty_lines(3)
add_body('[Student Name 1] — [Roll Number]')
add_body('[Student Name 2] — [Roll Number]')
add_body('[Student Name 3] — [Roll Number]')
add_body('[Student Name 4] — [Roll Number]')
add_empty_lines(2)
add_body('Date:')
add_body('Place: Delhi')

# ═══════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ═══════════════════════════════════════════════════════════
add_page_break()
add_centered('ACKNOWLEDGEMENT', 16, True, 24, 18)

add_body('We would like to express our sincere gratitude to our project supervisor, [Supervisor Name], [Designation], Department of [Department Name], Delhi Technological University, for providing invaluable guidance, constructive feedback, and constant encouragement throughout the duration of this project. The insightful suggestions at every stage of development shaped the direction and quality of this work.', indent=True)

add_body('We extend our thanks to [HOD Name], Head of the Department of [Department Name], for providing the necessary infrastructure and academic environment that facilitated the completion of this project.', indent=True)

add_body('We are grateful to the faculty members of the Department of [Department Name] for the knowledge and skills imparted during the course of our B.Tech programme, which formed the foundation upon which this project was built.', indent=True)

add_body('We also thank our fellow students for their constructive discussions and the collegial atmosphere that contributed to a productive working environment.', indent=True)

add_body('Finally, we are deeply indebted to our families for their unwavering support, patience, and encouragement throughout our academic journey.', indent=True)

add_empty_lines(3)
add_body('[Student Name 1]')
add_body('[Student Name 2]')
add_body('[Student Name 3]')
add_body('[Student Name 4]')

# ═══════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════
add_page_break()
add_centered('ABSTRACT', 16, True, 24, 18)

add_body('The rapid adoption of the Unified Payments Interface (UPI) in India, processing over 11 billion transactions per month, has created a parallel increase in fraudulent activities targeting digital payment users. Traditional rule-based fraud detection systems suffer from high false positive rates and inability to adapt to evolving fraud tactics. This project presents the design, implementation, and evaluation of a real-time UPI fraud detection system that employs XGBoost gradient boosting classification combined with a comprehensive feature engineering pipeline tailored to UPI transaction characteristics.', indent=True)

add_body('The system transforms six raw transaction attributes into twenty-eight derived features spanning five categories: amount characteristics, balance dynamics, temporal patterns, user behavioural history, and categorical encodings. The XGBoost classifier, trained on 50,000 synthetic UPI transactions with SMOTE (Synthetic Minority Over-sampling Technique) applied to address class imbalance, achieves a fraud detection recall of 94 percent and a precision of 85 percent on the held-out test set, with an ROC-AUC score of 0.987.', indent=True)

add_body('The trained model is deployed as a FastAPI inference service with a prediction latency of 28 milliseconds. An Express.js backend orchestrates the complete transaction lifecycle including authentication, fraud scoring, balance management, and alert generation. A React-based dashboard provides real-time monitoring with KPI visualisations, transaction management, alert handling, and analytical views.', indent=True)

add_body('The system implements a dual-mode detection architecture where a rule-based fallback mechanism activates automatically when the machine learning service is unavailable, ensuring uninterrupted fraud monitoring. Comparative evaluation demonstrates that the machine learning model outperforms the rule-based baseline across all metrics, with a 22 percentage-point improvement in fraud recall and a 43 percentage-point improvement in fraud precision.', indent=True)

add_body('Keywords: UPI, fraud detection, machine learning, XGBoost, gradient boosting, feature engineering, SMOTE, real-time prediction, web application, React, FastAPI, Express.js', bold=True)


# ═══════════════════════════════════════════════════════════
# Helper: Parse and add markdown content from file
# ═══════════════════════════════════════════════════════════
def process_markdown_file(filepath):
    """Read a markdown file and add its content to the Word doc."""
    with open(filepath, 'r') as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_headers = []
    table_rows = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table processing
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]

            # Skip separator rows (|---|---|)
            if all(c.replace('-', '').replace(':', '') == '' for c in cells):
                i += 1
                continue

            if not in_table:
                in_table = True
                table_headers = cells
                table_rows = []
            else:
                table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table
            if table_headers and table_rows:
                add_table_from_data(table_headers, table_rows)
                doc.add_paragraph()  # spacing after table
            in_table = False
            table_headers = []
            table_rows = []

        # Headings (skip # CHAPTER lines — we add those manually with page breaks)
        if line.startswith('# CHAPTER'):
            # Add page break and chapter heading
            add_page_break()
            add_centered(line.lstrip('# '), 16, True, 24, 18)
            i += 1
            continue
        elif line.startswith('# '):
            # Top-level heading (non-chapter)
            add_heading_custom(line.lstrip('# '), level=1)
            i += 1
            continue
        elif line.startswith('### '):
            add_heading_custom(line.lstrip('# '), level=3)
            i += 1
            continue
        elif line.startswith('## '):
            add_heading_custom(line.lstrip('# '), level=2)
            i += 1
            continue

        # Empty lines
        if line.strip() == '':
            i += 1
            continue

        # Horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # List items
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip().lstrip('-* ').strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold markdown
            text = re.sub(r'`(.*?)`', r'\1', text)  # Remove code markdown
            p = doc.add_paragraph(text, style='List Bullet')
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold markdown
        text = re.sub(r'`(.*?)`', r'\1', text)  # Remove code markdown
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Remove links

        if text:
            add_body(text, indent=True)

        i += 1

    # Flush remaining table
    if in_table and table_headers and table_rows:
        add_table_from_data(table_headers, table_rows)


# ═══════════════════════════════════════════════════════════
# Process all chapter files
# ═══════════════════════════════════════════════════════════
report_dir = os.path.dirname(os.path.abspath(__file__))

chapter_files = [
    'chapter-1-introduction.md',
    'chapter-2-literature-survey.md',
    'chapter-3-methodology.md',
    'chapter-4-system-architecture.md',
    'chapter-5-database-design.md',
    'chapter-6-ml-workflow.md',
    'chapter-7-implementation.md',
    'chapter-8-results.md',
    'chapter-9-screenshots.md',
    'chapter-10-testing.md',
    'chapter-11-conclusion.md',
    'chapter-12-future-scope.md',
    'references.md',
]

for chapter_file in chapter_files:
    filepath = os.path.join(report_dir, chapter_file)
    if os.path.exists(filepath):
        print(f'Processing: {chapter_file}')
        process_markdown_file(filepath)
    else:
        print(f'WARNING: File not found: {filepath}')

# ═══════════════════════════════════════════════════════════
# Save the document
# ═══════════════════════════════════════════════════════════
output_path = os.path.join(report_dir, 'UPI_Fraud_Detection_Report.docx')
doc.save(output_path)
print(f'\nReport saved to: {output_path}')
print(f'Total paragraphs: {len(doc.paragraphs)}')
